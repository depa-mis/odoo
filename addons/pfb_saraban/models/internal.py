from odoo import models, fields, api, _
from odoo.exceptions import UserError, ValidationError
from datetime import timedelta, datetime, date
import logging
import json
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches
import requests as r
_logger = logging.getLogger(__name__)


# def send(msg):
#     token = 'GB3DOAbe7Eu5vvVUQ3UFMMcNE8No1YqwR7zj3Rp9hDE'
#     headers = {'content-type': 'application/x-www-form-urlencoded',
#                'Authorization': 'Bearer '+token}
#     url = 'https://notify-api.line.me/api/notify'
#     r.post(url, headers=headers, data={'message': msg})


class internalDocument(models.Model):
    _name = 'document.internal.main'
    _order = 'date_document desc'
    _description = 'หนังสือภายใน'
    _inherit = ['mail.thread', 'mail.activity.mixin']

    # @api.model
    # def get_needaction_count(self):
    #     """ compute the number of needaction of the current user """
    #
    #     _logger.error('Call to needaction_count without partner_id')
    #     return 10

    def _default_employee(self):
        return self.env['hr.employee'].search([('user_id', '=', self.env.uid)], limit=1)

    def _get_director(self):
        print('GET DIRECTOR')
        sign_employee = {}
        director = self.env['fw_pfb_fin_settings2'].sudo().search([])
        for directoroffice in director.directorOfOffice:
            sign_employee = directoroffice
        return sign_employee

    # @api.multi
    # @api.depends('setting_line_ids')
    # def _compute_menu_count(self):
    #     count = 0.0
    #     for rec in self:
    #         for line in rec.setting_line_ids:
    #             if line.status_approve == '0':
    #                 count = count +1
    #         rec.menu_count_document = count
    #         print(rec.menu_count_document)
    prefix=fields.Char(
        string='prefix',
        copy=True,
        store=True,
        related='department_id.depa_document_sequence',
    )
    name = fields.Char(string='Sequence', required=True, copy=False, readonly=True, states={
                       'draft': [('readonly', False)]}, index=True, default=lambda self: _('New'),track_visibility='onchange')
    name_real = fields.Char(
        string='Real name',
        copy=False,
        track_visibility='onchange',
        store=True,
    )
    document_note = fields.Text(
        string="Memo",
        required=False,
    )
    # check_real = fields.Char(string='Just for check real name', copy=False)
    active = fields.Boolean('Active', default=True,
                            help="If unchecked, it will allow you to hide without removing it.")
    # menu_count_document = fields.Integer(compute='_compute_menu_count',store=True)
    # name_issue = fields.Char(string='issue', copy=False, states={
    #                          'draft': [('readonly', False)]}, index=True,track_visibility='onchange')
    # year_issue = fields.Char(string='issue date', copy=False, states={
    #                          'draft': [('readonly', False)]}, index=True,track_visibility='onchange')
    # check_issue = fields.Char(string='Just for check issue name', copy=False)
    date_document_real = fields.Date(string='Real Date', copy=False,track_visibility='onchange')
    """ document_type and document_type_select are type of document that you don have to do anything about it"""
    document_type = fields.Selection([('บันทึกข้อความ', 'บันทึกข้อความ'),
                                      ('คำสั่ง ก', 'คำสั่ง ก'),
                                      ('คำสั่ง ข', 'คำสั่ง ข'),
                                      ('คำสั่ง ค', 'คำสั่ง ค'),
                                      ('คำสั่ง พ', 'คำสั่ง พ'),
                                      ('ประกาศ', 'ประกาศ'),
                                      ('ประกาศพัสดุ', 'ประกาศพัสดุ'),
                                      ('ข้อบังคับ', 'ข้อบังคับ'),
                                      ('หนังสือภายนอก+หนังสือรับรอง',
                                       'หนังสือภายนอก+หนังสือรับรอง'),
                                      ('ระเบียบ', 'ระเบียบ'),
                                      ('หนังสือรับรอง', 'หนังสือรับรอง'),
                                      ], string='Document type',track_visibility='onchange',)
    document_type_select = fields.Selection([
        ('คำสั่ง ก', 'คำสั่ง ก'),
        ('คำสั่ง ข', 'คำสั่ง ข'),
        ('คำสั่ง ค', 'คำสั่ง ค'),
        ('คำสั่ง พ', 'คำสั่ง พ'),
        ('ประกาศ', 'ประกาศ'),
        ('ประกาศพัสดุ', 'ประกาศพัสดุ'),
        ('ข้อบังคับ', 'ข้อบังคับ'),
        ('หนังสือภายนอก+หนังสือรับรอง', 'หนังสือภายนอก+หนังสือรับรอง'),
        ('ระเบียบ', 'ระเบียบ'),
        ('หนังสือรับรอง', 'หนังสือรับรอง'),
    ], string='Document type select', default='หนังสือภายนอก+หนังสือรับรอง')

    date_document = fields.Date(
        string='Date',
        readonly=True,
        copy=False,
        default=lambda self: date.today(),
        store=True,
    )
    circular_letter = fields.Boolean(
        string='หนังสือเวียน',
        default=False,
        track_visibility='onchange'
    )
    state = fields.Selection([('cancel', 'Reject'),
                              ('draft', 'Draft'),
                              ('sent', 'Sent'),
                              ('1', 'ผอ.ฝ่าย'),
                              ('2', 'ฝ่ายที่เกี่ยวข้อง-1'),
                              ('3', 'เลขา ชสศด./รสศด.'),
                              ('4', 'ชสศด.'),
                              ('5', 'รสศด.'),
                              ('6', 'ฝ่ายที่เกี่ยวข้อง-2'),
                              ('7', 'ผสศด.'),
                              ('done', 'เสร็จสมบูรณ์'),
                              ],
                             readonly=True,
                             default='draft',
                             copy=False,
                             string="Status",
                             track_visibility='onchange')
    # date_document=fields.Datetime(string='Date',readonly=True,copy=False)
    # department_id = fields.Many2one('hr.department', string='Department', readonly=True)
    employee_id = fields.Many2one('hr.employee', string="Employee", default=_default_employee,
                                  required=True, ondelete='cascade', index=True, copy=False)
    department_id = fields.Many2one('hr.department', string="Department id",
                                    related="employee_id.department_id", readonly=True, copy=False)
    department_name = fields.Char(
        string="Department", related="employee_id.department_id.name", readonly=True, copy=False, store=True)
    speed = fields.Selection([('1', 'Normal'),
                              ('2', 'Express'),
                              ('3', 'urgent'),
                              ('4', 'The most urgent'),
                              ],
                             string="Speed ​​class", default='1')
    secret = fields.Selection([('normal', 'Normal'),
                               ('secret', 'Secret'),
                               ('vsecret', 'very secret'),
                               ('vvurgent', 'The most secret'),
                               ],
                              string="Secret ​​class",
                              default='normal')

    dear = fields.Selection([
        ('empolyee_name', 'empolyee name'),
        ('job_name', 'Job position')], string="Dear", default='empolyee_name')
    dear_select = fields.Char(
        string='เรียน',
    )
    # dear_select = fields.Selection([('เรียน', 'เรียน'),
    #                                 ('กราบเรียน', 'กราบเรียน'),
    #                                 ('กราบทูล', 'กราบทูล'),
    #                                 ('ขอประทานกราบทูล', 'ขอประทานกราบทูล'),
    #                                 ('ทูล', 'ทูล'),
    #                                 ('นมัสการ (สมเด็จและรองสมเด็จราชาคณะ)',
    #                                  'นมัสการ (สมเด็จและรองสมเด็จราชาคณะ)'),
    #                                 ('นมัสการ (พระราชาคณะ)',
    #                                  'นมัสการ (พระราชาคณะ)'),
    #                                 ('นมัสการ (พระภิกษุสงฆ์)',
    #                                  'นมัสการ (พระภิกษุสงฆ์)'),
    #                                 ])
    # dear_select_text = fields.One2many('document.internal.main.dear',
    #                                    'setting_id', string='dear', copy=True)
    job_name = fields.Many2many('hr.job', string="Job name")
    empolyee_name = fields.Many2many('hr.employee', string="Employee name")
    subject = fields.Char(string='Subject',track_visibility='onchange')
    sign = fields.Many2one(
        'hr.job',
        string="Sign"
    )
    sign_employee = fields.Many2one(
        'hr.employee',
        string="Sign",
        default=_get_director,
        store=True,
    )
    sign_department = fields.Many2one(
        'hr.job',
        string='Sign Job',
        related='sign_employee.job_id',
    )
    selected_material= fields.Many2one('document.template.setting', string='Document template')
    material = fields.Text(string='Material')
    for_document = fields.Selection(
        [('เพื่อโปรดดำเนินการ', 'เพื่อโปรดดำเนินการ'),
         ('เพื่อโปรดพิจารณา', 'เพื่อโปรดพิจารณา'),
         ('เพื่อโปรดทราบ', 'เพื่อโปรดทราบ'),
         ('เพื่อโปรดเห็นชอบ', 'เพื่อโปรดเห็นชอบ'),
         ('เพื่อโปรดมอบหมาย', 'เพื่อโปรดมอบหมาย'),
         ('เพื่อโปรดลงนาม', 'เพื่อโปรดลงนาม'),
         ('เพื่อโปรดอนุมัติ', 'เพื่อโปรดอนุมัติ'),
         ('เพื่อโปรดถือปฏิบัติ', 'เพื่อโปรดถือปฏิบัติ'),
         ('other', 'อื่นๆ'),],
        default=False
    )
    for_document_other = fields.Char(string="Other")
    note = fields.Text(string="Note",)
    routings_internal_id = fields.Many2one(
        'document.internal.setting', string='Document routing', copy=True, domain="[ ('is_active', '=', True)]")
    setting_line_ids = fields.One2many(
        'document.internal.main.setting.line',
        'setting_id',
        string='Document main line',
        copy=True,
        require=True,
    )
    waiting_line_ids = fields.One2many(
        'waiting.document.main.setting.line',
        'setting_id',
        string='Document wating line',
        copy=False,
        require=True
    )
    reference_line_ids_multi = fields.Many2many(
        comodel_name="document.internal.main",
        relation="document_internal_rel",
        column1="track",
        culomn2="chair",
        string="Reference Line",
        domain="[('state','!=','draft')]"
    )
    reference_receive_document = fields.Many2many(
        'receive.document.main',
        string='reference receive book',
        copy=True,
        domain="[('state','!=','draft')]"
    )
    reference_receive_document_related = fields.Many2many(
        'receive.document.main',
        string='reference receive book',
        copy=True,
        readonly=1,
        related='reference_receive_document',
    )
    # reference_line_ids=fields.Many2many('document.internal.main','document.internal.main',relation="name",column_1='document_type',string="reference internal book")
    approval_count = fields.Integer(
        string="Approval Count",
        default=0,
        copy=False
    )
    check_group_user = fields.Boolean(
        string="Invisible",
        compute="_check_group_user",
        store=False
    )
    attachment_ids = fields.Many2many(
        'ir.attachment',
        "attachment_document_receive_rel",
        "attachment_document_receive_id",
        "attachment_id",
        string='Files',
        track_visibility='onchange'
    )
    # attachment_log = fields.Text(string='attacment log', store=True)
    # allow_group=fields.Many2many(
    #     'hr.employee',
    #     string="Allow name"
    # )
    send_with_original_document = fields.Boolean(
        string='Sent with original document',
        default=False,
    )
    internal_title = fields.Char(
        string='Title',
    )
    setting_line_ids_related = fields.One2many(
        'document.internal.main.setting.line',
        'setting_id',
        related='setting_line_ids',
        string='Document main line',
        copy=False,
        require=True
    )
    attached_file_in = fields.Many2many(
        'ir.attachment',
        "attachment_file_in_document_receive_rel",
        "attachment_file_in_document_receive_id",
        "attachment_id",
        string='Attached file',
        track_visibility='onchange'
    )
    is_participant = fields.Boolean(
        string='Is participant',
        compute='_check_is_participant',
        default=True,
    )
    approver_name_list = fields.Text(
        string='Approver Name',
        compute='_compute_setting_line',
        store=True
    )
    sign_email = fields.Char(
        string='Email'
    )
    suffix = fields.Char(
        string='คำลงท้าย'
    )

    fin_100_ids = fields.Many2many(
        'fw_pfb_fin_system_100',
        'fin100_internal_document_rel',
        'fin_internal_document_id',
        'fin_id',
        string='FIN100'
    )

    fin_201_ids = fields.Many2many(
        'fw_pfb_fin_system_201',
        'fin201_internal_document_rel',
        'fin_internal_document_id',
        'fin_id',
        string='FIN201'
    )
    
    fin_401_ids = fields.Many2many(
        'fw_pfb_fin_system_401',
        'fin401_internal_document_rel',
        'fin_internal_document_id',
        'fin_id',
        string='FIN401'
    )
    tag_ids=fields.Many2many(
        'document.internal.document.setting.tag',
        'internal_setting_tag_rel',
        'internal_id',
        'setting_id',
        string="tag",
        track_visibility='onchange',
        copy=False
    )
    internal_document_note = fields.Text(
        string='Note'
    )
            
    @api.multi
    def _setting_line_approver_set(self):
        internal_obj = self.env['document.internal.main'].search([])
        for rec in internal_obj:
            approver_list = ''
            if rec.setting_line_ids:
                for sti in rec.setting_line_ids:
                    if sti.employee_id.name:
                        approver_list += sti.employee_id.name + ' '
                rec.approver_name_list = approver_list

    @api.multi
    @api.depends('setting_line_ids')
    def _compute_setting_line(self):
        for rec in self:
            approver_list = ''
            if rec.setting_line_ids:
                for sti in rec.setting_line_ids:
                    if sti.employee_id.name:
                        approver_list += sti.employee_id.name + ' '
                rec.approver_name_list = approver_list

    @api.multi
    def _check_is_participant(self):
        for rec in self:
            sli_ids = []
            for sli in rec.setting_line_ids:
                sli_ids.append(sli.employee_id.user_id.id)
            if self._uid == rec.employee_id.user_id.id or self._uid in sli_ids:
                rec.is_participant = True
            else:
                rec.is_participant = False

    # _sql_constraints = [
    #     ('check_real', 'unique(check_real)', "ชื่อ ฉบับไม่ควรซำ"),
    # ]
    # def archived(self):
    #     self.active= not self.active
        # raise ValidationError(_(self.active))

        
    #history log
    # @api.onchange('document_type')
    # def __change_document_type(self):
    #     if self.document_type:
    #         self.message_post(body=_("เปลียนประเภทเอกสารเป็น"))
    #         _logger.info("เปลียนประเภทเอกสารเป็น test")
    
    # @api.onchange('document_type')
    # def __change_document_type(self):
    #     self.message_post(body="เปลียนประเภทเอกสารเป็น"+self.document_type)

    @api.model
    def create(self, vals):
        if vals.get('name', _('New')) == _('New'):
            vals['name'] = self.env['ir.sequence'].next_by_code('esaraban.internal') or _('New')
            # vals['date_document']=fields.Date.today()
        result = super(internalDocument, self).create(vals)
        return result
            
    def make_order_setting_line(self):
        setting_line_ids = self.setting_line_ids.sorted(lambda x: x.step)
        sequence = 0
        for slid in setting_line_ids:
            slid.update({"sequence":sequence})
            sequence += 1

    @api.depends('employee_id')
    def _check_group_user(self):
        _logger.info('group_super_user_document =============== %s' %
                     self.env.user.has_group('pfb_saraban.group_super_user_document'))
        if self.env.user.has_group('pfb_saraban.group_super_user_document'):
            _logger.info('IFFFFFF')
            self.update({
                'check_group_user': True
            })
        else:
            _logger.info('ELSEEEE')
            self.update({
                'check_group_user': False
            })
    # @api.onchange('setting_line_ids')
    # def _onchange_waiting(self):
        
    # @api.onchange('message_main_attachment_id')
    # def test(self):
    #     self.message_post(body="test")
    # @api.onchange('attachment_ids')
    # def _compute_attachments(self):
    #     try:
    #         # body="abcd"
    #         # self.send_message(body)
    #         self.message_post(body="test")
            # temp = self.attachment_log
            # # raise ValidationError(_(self.attachment_log))
            # # send(" "+self.attachment_log)
            # old_attac = []
            # if self.attachment_log:
            #     old_attac = json.loads(temp)
            # attctment_list = []
            # for attcment in self.attachment_ids:
            #     attctment_list.append(attcment.name)

    #         self.attachment_log = json.dumps(attctment_list, sort_keys=True)
    #         diff_name = list(list(set(old_attac)-set(attctment_list)) +
    #                          list(set(attctment_list)-set(old_attac)))
    #         # diff_name=[new_file_name for new_file_name in li1 + li2 if new_file_name not in li1 or new_file_name not in li2]
    #         # raise ValidationError(_(diff_name))
    #         str = ""
    #         for names in diff_name:
    #             str += names+" "
    #         # raise ValidationError(_(self.attachment_log))
    #         send(str)

    #         # send('old '+temp)
    #         # if self.attachment_log !="[]":
    #         #     body="abcd"
    #         #     self.message_post(body=_(body))
        # except Exception as e:
        #     raise ValidationError(_(e))

    # def send_message(self, body):
    #     self.message_post(body=_(body))
    # @api.onchange('attachment_ids')
    # def _test_onchange(self):
    #     self.message_post(body="Sent")

    # @api.onchange('name_real')
    # def onchange_real_name(self):
    #     try:
    #         if self.name_real:
    #             # raise ValidationError(_('Condition not working'))
    #             if self.document_type and self.name_real:
    #                 self.check_real = str(self.document_type)+self.department_name+self.name_real
    #     except Exception as e:
    #         raise ValidationError(_(e))

    @api.onchange('selected_material')
    def onchange_selected_material(self):
        self.ensure_one()
        if self.material:
            self.material+=self.selected_material.material
        else:
            self.material=self.selected_material.material
        
    @api.onchange('routings_internal_id')
    def onchange_routings_internal_id(self):
        self.ensure_one()
        self.update({
            'setting_line_ids': False
        })
        main_line_ids = self.routings_internal_id.user_id
        values = []
        employee_ids = []
        for line_id in main_line_ids:
            employee_ids.append(line_id.employee_id.id)
            # self.empolyee_name = [(4, line_id.employee_id.id)]
            
            values.append([0, 0, {
                'employee_id': line_id.employee_id.id,
                'job_id_name': line_id.job_id_name,
                'status': line_id.status,
                'step': line_id.step,
                'approve_type': line_id.approve_type,
                'is_active': line_id.is_active,
            }])
        self.update({
            'setting_line_ids': values,
            'empolyee_name': [(6, _, employee_ids)],
        })

    @api.onchange('document_type_select')
    def onchange_document_type_select(self):
        for record in self:
            if record.document_type_select:
                self.update({'document_type': record.document_type_select})
        # raise ValidationError(_('Dummy not checked'))

    def action_make_approval_wizard(self):
        self.ensure_one()
        approve_type = False
        employee_id = False
        setting_line = False
        status = False
        count_nonactive = 0
        for line in self.setting_line_ids:
            if not line.is_active:
                count_nonactive += 1
        if self.setting_line_ids:
            for line in self.setting_line_ids:
                if not line.approve_time and line.is_active and self.env.uid == line.employee_id.user_id.id and line.status_approve == '0':
                    approve_type = line.approve_type
                    employee_id = line.employee_id
                    setting_line = line.id
                    status = line.status
                    step = line.step
                    return {
                        'name': "Make Approval Wizard",
                        'view_mode': 'form',
                        'view_id': False,
                        'view_type': 'form',
                        'res_model': 'make.approval.wizard',
                        'type': 'ir.actions.act_window',
                        'target': 'new',
                        'context': {
                            'default_approve_type': approve_type,
                            'default_employee_id': employee_id.id,
                            'default_setting_line': setting_line,
                            'default_status': status,
                            'default_total_for_approve': len(self.setting_line_ids)-count_nonactive,
                            'default_setting_id': self.id,
                        }
                    }
            raise ValidationError(_('Not your turn'))

    def action_sent_to_supervisor(self):
        # if self.employee_id.dummy:
        #     raise ValidationError(_('Dummy not checked'))
        self.ensure_one()
        sequence_id = False
        
        for line in self.setting_line_ids:
            if line.dummy and line.is_active:
                raise ValidationError(_('Dummy not checked'))
        if(len(self.setting_line_ids) == 0):
            raise ValidationError(_('Please add the line'))

        # ISSUE 224 ถอย
        if self.document_type == 'บันทึกข้อความ':
            if self.name_real in ['', ' ',None,False]:
                sequence_id = self.env['ir.sequence'].search([
                    ('department_id', '=', self.department_id.id),
                    ('document_type', '=', self.document_type),
                    ('sarabun_active', '=', True),
                ]) or False
                if sequence_id:
                    self.update({
                        'date_document_real': fields.Date.today(),
                    })
                    self.name_real = sequence_id.next_by_id()
                else:
                    raise ValidationError(_('Please add the sequence'))

        step_array = []
        for line in self.setting_line_ids:
            if line.is_active:
                step_array.append(line.step)
        initial_step = min(step_array)

        for line in self.setting_line_ids:
            if line.step == initial_step and line.is_active:
                line.status_approve = '0'
        # raise ValidationError(_(fields.Date.today()))
        self.update({
            'state': 'sent',
            
        })
        # self.allow_group=False
        # for setting_line_id in self.setting_line_ids:
        #     values = []
        #     if setting_line_id.status_approve=='0':
        #         values.append([0, 0, {
        #             'allow_group': setting_line_id.employee_id
        #         }])
        #     _logger.info("name="+setting_line_id.employee_id.name)
        
        # for i in self.empolyee_name:
        #     print('empolyee_name=')
        #     print(i)
        
        # self.update({
        #     'allow_group': values
        # })
        self.update({
            'waiting_line_ids': False
        })
        main_line_ids = self.setting_line_ids
        values = []
        for line_id in main_line_ids:
            if line_id.is_active and line_id.status_approve=='0':
                values.append([0, 0, {
                    'employee_id': line_id.employee_id.id
                }])
        self.update({
            'waiting_line_ids': values
        })
        self.message_post(body="Sent")
    
    def set_to_draft(self):
        self.update({
            'state': 'draft',
            'approve_count': 0,
        })
        self.message_post(body="Set to draft")
        for line in self.setting_line_ids:
            line.status_approve = ''
            line.comment = ''
            line.approve_time = ''

    def state_done(self):
        sequence_id = self.env['ir.sequence'].search([
            ('document_type', '=', 'หนังสือภายนอก+หนังสือรับรอง'),
        ]) or False
        if sequence_id:
            self.update({
                'date_document_real': fields.Date.today(),
            })
        self.update({
            'name': sequence_id.next_by_id(),
            'state': 'done'
        })

    @api.model
    def _cron_send_email(self):
        users = self.search(
            [('setting_line_ids.employee_id.name', '=', 'Admin1')])
        for user in users:
            digests = self.search(
                [('setting_line_ids.employee_id.name', '=', 'Admin1')])
            for digest in digests:
                _logger.warning(digest.name)
                
            self.env['document.internal.main.setting.line'].search([
                    ('status_approve', '=', '0'),
                    ('employee_id', '=', self.employee_id.id)])
            # _logger.info(digest.name)
            # try:
            #     digest.action_send()
            # except MailDeliveryException as e:
            #     _logger.warning('MailDeliveryException while sending digest %d. Digest is now scheduled for next cron update.')
    # def action_send(self):
    #     for digest in self:
    #         for user in digest.user_ids:
    #             subject = '%s: %s' % (user.company_id.name, digest.name)
    #             digest.template_id.with_context(user=user).send_mail(digest.id, force_send=True, raise_exception=True, email_values={'email_to': user.email, 'subject': subject})
    #         digest.next_run_date = digest._get_next_run_date()

    attachment_id = fields.Many2one('ir.attachment', 'Attachment')
    # setting_id = fields.Many2one('document.internal.main', string="document main")

    def from_data(self,current_id):
        rec_ids = self.env['document.internal.main'].search([('id','=',current_id)])
        document = Document()
        document.add_heading('Document Internal Main',4)
        document.add_paragraph('')
        for doc_ids in rec_ids:
            circular_letter = str(doc_ids.circular_letter)
            name = str(doc_ids.name)
            date = str(doc_ids.date_document)
            real_name = str(doc_ids.name_real)
            real_date = str(doc_ids.date_document_real)
            depart = str(doc_ids.department_name)
            speed = str(doc_ids.speed)
            secret = str(doc_ids.secret)
            subject = str(doc_ids.subject)
            dear = str(doc_ids.dear)
            employee = str(doc_ids.empolyee_name.name)
            sign = str(doc_ids.sign.name)
            for_document = str(doc_ids.for_document)
            other = str(doc_ids.for_document_other)
            note = str(doc_ids.note)
            material = str(doc_ids.material)
            
            fist_para = document.add_paragraph('หนังสือเวียน ')
            fist_para.add_run(circular_letter)
            sec_para = document.add_paragraph('Name ')
            sec_para.add_run(name)
            thir_para = document.add_paragraph('Date ')
            thir_para.add_run(date)
            four_para = document.add_paragraph('Real Name ')
            four_para.add_run(real_name)
            fift_para = document.add_paragraph('Real Date ')
            fift_para.add_run(real_date)
            six_para = document.add_paragraph('Department ')
            six_para.add_run(depart)
            sev_para = document.add_paragraph('Speed ')
            sev_para.add_run(speed)
            eigh_para = document.add_paragraph('Secret ')
            eigh_para.add_run(secret)
            nine_para = document.add_paragraph('Subject ')
            nine_para.add_run(subject )
            ten_para = document.add_paragraph('Dear ')
            ten_para.add_run(dear)
            sixteen_para = document.add_paragraph('Employee ')
            sixteen_para.add_run(employee)
            ele_para = document.add_paragraph('Sign ')
            ele_para.add_run(sign)
            twel_para = document.add_paragraph('For Document ')
            twel_para.add_run(for_document)
            thirteen_para = document.add_paragraph('Material ')
            thirteen_para.add_run(material)
            fourteen_para = document.add_paragraph('Other ')
            fourteen_para.add_run(other)
            fifteen_para = document.add_paragraph('Note ')
            fifteen_para.add_run(note)
            document.add_paragraph('')

#         document.save('test.docx')
        fp = BytesIO()
        document.save(fp)
        fp.seek(0)
        data = fp.read()
        fp.close()
        return data
    
   
    @api.multi
    def unlink(self):
        if self.state != 'draft':
            raise UserError(
                _('You cannot delete an document which is not draft.'))
        return super(internalDocument, self).unlink()
    
class InternalDocumentSettingLine(models.Model):
    _name = 'document.internal.main.setting.line'
    _description = 'สารบรรณ หนังสือภายใน บันทึกข้อความ'
    _order = 'sequence,id'

    setting_id = fields.Many2one(
        'document.internal.main', string="document main")
    employee_id = fields.Many2one(
        'hr.employee',
        string="Employee Name",
        index=True,
    )
    job_id_name = fields.Many2one(
        'hr.job',
        string="Job name"
    )
    """Step and state are the same so don have to do anything about it"""
    step = fields.Selection(
        string="Step",
        selection=[
            ('1', '1'),
            ('2', '2'),
            ('3', '3'),
            ('4', '4'),
            ('5', '5'),
            ('6', '6'),
            ('7', '7')
        ],
        default='1'
    )
    status = fields.Selection(
        string="Status",
        selection=[
        ('1', 'ผอ.ฝ่าย'),
        ('2', 'ฝ่ายที่เกี่ยวข้อง-1'),
        ('3', 'เลขา ชสศด./รสศด.'),
        ('4', 'ชสศด.'),
        ('5', 'รสศด.'),
        ('6', 'ฝ่ายที่เกี่ยวข้อง-2'),
        ('7', 'ผสศด.')
        ],
        readonly=True,
        compute='set_step'
    )
    dummy = fields.Boolean(
        string="dummy",
        related='employee_id.dummy',
        readonly=True
    )
    approve_type = fields.Selection(
        string="Approve type", selection=[
        ('require', 'Is require to approve'),
        ('comments', 'Comments only')]
    )
    approve_request_time = fields.Datetime(
        string='วันที่ต้องอนุมัติ/รับทราบ',
        copy=False
    )
    approve_time = fields.Datetime(
        string='approve time',
        copy=False
    )
    comment = fields.Html(
        string='comment',
        copy=False
    )
    """0=waiting 1=approve 2=approve 3=reject"""
    status_approve = fields.Selection(
        string="status",
        selection=[
        ('0', 'รออนุมัติ'),
        ('1', 'เห็นชอบ/อนุมัติแล้ว'),
        ('2', 'รับทราบแล้ว'),
        ('3', 'ไม่อนุมัติแล้ว')],
        copy=False
    )
    document_id = fields.Many2one(
        'document.internal.setting',
        string="Document id"
    )
    is_active = fields.Boolean(
        string='Active',
        default=True,
    )
    sequence = fields.Integer(
        string='Sequence',
        index=True
    )
    sender = fields.Many2one(
        'hr.employee',
        string="Sender",
        readonly=True,
    )

    # Backup field
    step_before_change = fields.Selection(
        help_text="Step before change from 6 to 7",
        string="Step",
        selection=[
            ('', ''),
            ('1', '1'),
            ('2', '2'),
            ('3', '3'),
            ('4', '4'),
            ('5', '5'),
            ('6', '6'),
            ('7', '7')
        ],
        default='',
    )
    max_date_approve = fields.Date(string='Step Date Approve', copy=False)

    @api.depends('step')
    def set_step(self):
        for record in self:
            if record.step:
                step = record.step
                record.update({'status': step})


class InternalDear(models.Model):
    _name = 'document.internal.main.dear'
    _description = 'เรียน'
    setting_id = fields.Many2one('document.internal.main')
    name = fields.Char(string="name")
    sequence = fields.Integer(string='Sequence', index=True)


class InternalDocumentReferenceLine(models.Model):
    _name = 'document.internal.main.reference.line'
    _description = 'สารบรรณ หนังสือภายใน'
    setting_id = fields.Many2one('document.internal.main')
    name_id = fields.Many2one('document.internal.main', string="name")
    name=fields.Char(
        string="Name text", related='name_id.name', readonly=True)
    # document_type= fields.Char(string="Document type", related='name_id.document_type',readonly=True)
    document_type = fields.Selection([('บันทึกข้อความ', 'บันทึกข้อความ'),
                                      ('คำสั่ง ก', 'คำสั่ง ก'),
                                      ('คำสั่ง ข', 'คำสั่ง ข'),
                                      ('คำสั่ง ค', 'คำสั่ง ค'),
                                      ('คำสั่ง พ', 'คำสั่ง พ'),
                                      ('ประกาศ', 'ประกาศ'),
                                      ('ข้อบังคับ', 'ข้อบังคับ'),
                                      ('หนังสือภายนอก+หนังสือรับรอง',
                                       'หนังสือภายนอก+หนังสือรับรอง'),
                                      ('ระเบียบ', 'ระเบียบ'),
                                      ], string='Document type', related='name_id.document_type', readonly=True)
    department_name = fields.Char(
        string="Department", related='name_id.department_name', readonly=True)
    subject = fields.Char(
        string='Subject', related='name_id.subject', readonly=True)



class WaitingDocumentSettingLine(models.Model):
    _name = 'waiting.document.main.setting.line'
    _description = 'สารบรรณ หนังสือภายใน บันทึกข้อความ'

    setting_id = fields.Many2one(
        'document.internal.main', string="document main")
    employee_id = fields.Many2one(
        'hr.employee',
        string="Employee Name"
    )
    sequence = fields.Integer(string='Sequence', index=True)
