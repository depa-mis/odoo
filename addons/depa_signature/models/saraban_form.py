import time
from datetime import datetime, timedelta
from dateutil import relativedelta
from odoo import models, fields, api, _
import base64
from io import BytesIO
from odoo.exceptions import UserError, ValidationError
from odoo.tools import DEFAULT_SERVER_DATE_FORMAT
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt, Mm, Inches, Length, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

from docxtpl import DocxTemplate

from dateutil.relativedelta import relativedelta
import os


class saraban_form(models.Model):
    _name = 'saraban_form'

    def year_convert(self, convert_date):
        date_converted = convert_date + relativedelta(years=543)
        date_converted = date_converted.strftime('%d/%m/%Y %H:%M:%S')
        return date_converted

    def year_convert2(self, convert_date):
        date_converted = convert_date + relativedelta(years=543)
        date_converted = date_converted.strftime('%d/%m/%Y')
        return date_converted

    def to_thai_number(self, number):
        thai_number = []
        for i in number:
            if i == '0':
                i = '๐'
            if i == '1':
                i = '๑'
            elif i == '2':
                i = '๒'
            elif i == '3':
                i = '๓'
            elif i == '4':
                i = '๔'
            elif i == '5':
                i = '๕'
            elif i == '6':
                i = '๖'
            elif i == '7':
                i = '๗'
            elif i == '8':
                i = '๘'
            elif i == '9':
                i = '๙'

            thai_number.append(i)

        return ''.join(thai_number)

    def to_single_thai_number(self, index):
        numbers = [
            '๐',
            '๑',
            '๒',
            '๓',
            '๔',
            '๕',
            '๖',
            '๗',
            '๘',
            '๙'
        ]
        if not index:
            return None

        return numbers[index]

    def text_speed(self, speed):
        thai_speed = ''
        if speed == 'normal':
            thai_speed = 'ปกติ'
        if speed == '2':
            thai_speed = 'ด่วน'
        if speed == '3':
            thai_speed = 'ด่วนมาก'
        if speed == '4':
            thai_speed = 'ด่วนที่สุด'
        return ''.join(thai_speed)

    def text_secret(self, speed):
        text_secret = ''
        if speed == 'normal':
            text_secret = 'ปกติ'
        if speed == 'secret':
            text_secret = 'ลับ'
        if speed == 'vsecret':
            text_secret = 'ลับมาก'
        if speed == 'vurgent':
            text_secret = 'ลับที่สุด'
        return ''.join(text_secret)

    def to_full_thai_date(self, orgin_date):
        day = self.to_thai_number(str(int((orgin_date).strftime('%d'))))
        month = (int((orgin_date).strftime('%m')))
        year = self.to_thai_number(str(int((orgin_date).strftime('%Y')) + 543))

        if month == 1:
            month_th = 'มกราคม'
        if month == 2:
            month_th = 'กุมภาพันธ์'
        if month == 3:
            month_th = 'มีนาคม'
        if month == 4:
            month_th = 'เมษายน'
        if month == 5:
            month_th = 'พฤษภาคม'
        if month == 6:
            month_th = 'มิถุนายน'
        if month == 7:
            month_th = 'กรกฎาคม'
        if month == 8:
            month_th = 'สิงหาคม'
        if month == 9:
            month_th = 'กันยายน'
        if month == 10:
            month_th = 'ตุลาคม'
        if month == 11:
            month_th = 'พฤศจิกายน'
        if month == 12:
            month_th = 'ธันวาคม'

        return f"{day} {month_th} {year}"

    # def mid(self, name_real):
    #     print(name_real[int(len(name_real) / 2) - 1:int(len(name_real) / 2) + 5])
    #     print(name_real[int(len(name_real) / 2) - 1:int(len(name_real) / 2) + 6])
    #     print(name_real[int(len(name_real) / 2) - 1:int(len(name_real) / 2) + 7])
    #     print(name_real[int(len(name_real) / 2) - 1:int(len(name_real) / 2) + 8])
    #     print(name_real[int(len(name_real) / 2) - 1:int(len(name_real) / 2) + 9])
    #     return name_real[int(len(name_real) / 2) - 1:int(len(name_real) / 2) + 2]

    def from_data(self, current_id):
        cur_path = os.path.dirname(os.path.abspath(__file__))
        docx_path = cur_path + '/template/depa_template3.docx'
        rec_ids = self.env['document.internal.main'].search([('id', 'in', [current_id])])
        # document = Document()
        document = DocxTemplate(docx_path)
        style = document.styles['Normal']
        font = style.font
        font.name = 'TH SarabunPSK'
        font.size = Pt(16)
        section = document.sections[0]
        footer = section.footer
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Inches(1.181102)
        section.right_margin = Inches(0.7874016)
        section.top_margin = Inches(0.3937007)
        section.bottom_margin = Inches(0.3937007)
        # section.top_margin = Inches(0.9842520)
        # section.bottom_margin = Inches(0.5905512)
        date = ''
        real_date = ''
        depa_department = ''
        depa_phone = ''
        subject = ''
        for_document = ''
        for_document_other = ''
        material = ''
        note = ''
        job_title = ''
        real_name = ''
        department_phone = ''
        department_fax = ''
        department_email = ''
        employee = []
        state_line2 = []
        dear_name = []
        # styles = document.styles
        # paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
        # for style in paragraph_styles:
        #     print(style.name)

        for doc_ids in rec_ids:
            if doc_ids.document_type == 'คำสั่ง ก' or doc_ids.document_type == 'คำสั่ง ข' or \
                    doc_ids.document_type == 'คำสั่ง ค' or doc_ids.document_type == 'คำสั่ง พ':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                speed = self.text_speed(str(doc_ids.speed))
                secret = self.text_secret(str(doc_ids.secret))
                print(secret, speed)
                sing_department = self.to_thai_number(str(doc_ids.sign_department.name))
                sing_employee = self.to_thai_number(str(doc_ids.sign_employee.name))
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.subject:
                    subject = self.to_thai_number(str(doc_ids.subject))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                if doc_ids.material:
                    material = str(doc_ids.material)
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                    day = self.to_thai_number(str(int((doc_ids.date_document).strftime('%d'))))
                    month = (int((doc_ids.date_document).strftime('%m')))
                    if month == 1:
                        month_th = 'มกราคม'
                    if month == 2:
                        month_th = 'กุมภาพันธ์'
                    if month == 3:
                        month_th = 'มีนาคม'
                    if month == 4:
                        month_th = 'เมษายน'
                    if month == 5:
                        month_th = 'พฤษภาคม'
                    if month == 6:
                        month_th = 'มิถุนายน'
                    if month == 7:
                        month_th = 'กรกฎาคม'
                    if month == 8:
                        month_th = 'สิงหาคม'
                    if month == 9:
                        month_th = 'กันยายน'
                    if month == 10:
                        month_th = 'ตุลาคม'
                    if month == 11:
                        month_th = 'พฤศจิกายน'
                    if month == 12:
                        month_th = 'ธันวาคม'
                    year = self.to_thai_number(str(int((doc_ids.date_document).strftime('%Y')) + 543))
                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                    len_real_name = str(len(doc_ids.name_real))
                for employee_name in doc_ids.empolyee_name:
                    employee.append(str(employee_name.name))
                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)
                        employee_line = str(doc_line.employee_id.name)
                        employee_job = str(doc_line.job_id_name.name)
                department_obj = self.env['hr.department'].search([
                    ('name', '=', doc_ids.department_name),
                ])
                for department_line in department_obj:
                    if department_line.depa_phone:
                        department_phone = department_line.depa_phone
                    if department_line.depa_fax:
                        department_fax = department_line.depa_fax
                    if department_line.depa_email:
                        department_email = department_line.depa_email

                document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER

                zero_record = document.add_paragraph('')
                # font1 = zero_record.add_run('คำสั่ง').font
                # font1.size = Pt(18)
                # font1.bold = True
                zero_record.add_run('')
                zero_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                zero_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                zero_record.paragraph_format.space_after = Pt(0)

                one_record = document.add_paragraph('')
                font1 = one_record.add_run('คำสั่งสำนักงานส่งเสริมเศรษฐกิจดิจิทัล').font
                font1.size = Pt(18)
                font1.bold = True
                one_record.add_run('')
                one_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                one_record.paragraph_format.space_after = Pt(0)

                two_record = document.add_paragraph('')
                font1 = two_record.add_run('ที่  ').font
                font2 = two_record.add_run(name).font
                font1.size = Pt(18)
                font1.bold = True
                font2.size = Pt(18)
                font2.bold = True
                two_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                two_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                two_record.paragraph_format.space_after = Pt(0)

                three_record = document.add_paragraph('')
                font1 = three_record.add_run('เรื่อง ')
                font2 = three_record.add_run(subject).font
                font1.size = Pt(18)
                font1.bold = True
                font2.size = Pt(18)
                font2.bold = True
                three_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                three_record.paragraph_format.space_after = Pt(0)
                three_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                five_record = document.add_paragraph('____________________________________  ')
                five_record.add_run('')
                five_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                edit_record = document.add_paragraph('')
                edit_record.add_run('')
                edit_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record.paragraph_format.space_after = Pt(0)
                edit_record.paragraph_format.space_before = Pt(12)

                six_record = document.add_paragraph('')
                six_record.add_run(material)
                six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                six_record.paragraph_format.space_after = Pt(0)
                six_record.paragraph_format.space_before = Pt(6)

                six_record2 = document.add_paragraph('ทั้งนี้ ตั้งแต่บัดนี้เป็นต้นไป')
                six_record2.add_run('')
                six_record2.paragraph_format.left_indent = Inches(0.9842520)
                six_record2.paragraph_format.space_after = Pt(0)
                six_record2.paragraph_format.space_before = Pt(6)

                seven_record = document.add_paragraph('สั่ง ณ วันที่ ')
                seven_record.add_run(day)
                seven_record.add_run(' ')
                seven_record.add_run('เดือน ')
                seven_record.add_run(month_th)
                seven_record.add_run(' ')
                seven_record.add_run('พ.ศ. ')
                seven_record.add_run(year)
                seven_record.paragraph_format.left_indent = Inches(1.574803)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record3 = document.add_paragraph('')
                edit_record3.add_run('')
                edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record3.paragraph_format.space_after = Pt(0)
                edit_record4 = document.add_paragraph('')
                edit_record4.add_run('')
                edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record4.paragraph_format.space_after = Pt(0)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(sing_employee)
                sevens_record.add_run(' )')
                sevens_record.paragraph_format.left_indent = Inches(3.149606)

                eight_record = document.add_paragraph('')
                eight_record.add_run(sing_department)
                eight_record.paragraph_format.left_indent = Inches(3.051181)
                edits_record = document.add_paragraph('')
                edits_record.add_run('')

                font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_job_title).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(current_date).font
                font.size = Pt(10)
            if doc_ids.document_type == 'บันทึกข้อความ':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.subject:
                    subject = self.to_thai_number(str(doc_ids.subject))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                if doc_ids.material:
                    material = str(doc_ids.material)
                state = str(doc_ids.state)
                sign = str(doc_ids.sign.name)
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if obj.job_title:
                    job_title = str(obj.job_title)
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                if doc_ids.date_document_real:
                    real_date = str(self.year_convert2(doc_ids.date_document_real))
                job_em = []
                for employee_name in doc_ids.setting_line_ids_related:
                    if employee_name.is_active:
                        employee.append(str(employee_name.employee_id.name))
                        job_em.append(str(employee_name.employee_id.job_title))
                print(employee_name, employee, job_em)

                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)

                document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                header_fist = document.add_paragraph('                                                 ').add_run(
                    'บันทึกข้อความ')
                header_fist.bold = True
                font = header_fist.font
                font.name = 'TH SarabunPSK'
                font.size = Pt(24)
                header_fist.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record = document.add_paragraph('ฝ่าย/ส่วน ')
                # paragraph_format = header_fist.paragraph_format
                # paragraph_format.line_spacing = Pt(18)
                one_record.underline = WD_UNDERLINE.DOUBLE
                one_record.add_run(depa_department)
                one_record.add_run('                                                             ')
                one_record.add_run('โทร. ')
                one_record.add_run('')
                one_record.add_run(depa_phone)
                two_record = document.add_paragraph('ที่ ')
                two_record.add_run(real_name)
                two_record.add_run(
                    '                                                   ')
                two_record.add_run('วันที่ ')
                two_record.add_run(date)
                three_record = document.add_paragraph('เรื่อง ')
                three_record.add_run(subject)
                four_record = document.add_paragraph('เรียน ')
                four_record.add_run(employee[-1])
                five_record = document.add_paragraph('ผ่าน ')
                for line_name2 in state_line2:
                    five_record.add_run(' ')
                    five_record.add_run(line_name2)
                    five_record.add_run(',')
                six_record = document.add_paragraph('')
                six_record.add_run(material)
                seven_record = document.add_paragraph('')
                seven_record = document.add_paragraph('จึงมาเรียน')
                seven_record.add_run(for_document)
                seven_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                seven_record.paragraph_format.left_indent = Inches(0.9842520)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(employee[-1])
                sevens_record.add_run(' )')
                sevens_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                sevens_record.paragraph_format.left_indent = Inches(3.444882)

                eight_record = document.add_paragraph('')
                eight_record.add_run(job_em[-1])
                eight_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                eight_record.paragraph_format.left_indent = Inches(3.346457)

                font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_job_title).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(current_date).font
                font.size = Pt(10)
            if doc_ids.document_type == 'หนังสือภายนอก+หนังสือรับรอง':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                speed = self.text_speed(str(doc_ids.speed))
                secret = self.text_secret(str(doc_ids.secret))
                print(doc_ids.secret)
                dear_select = str(doc_ids.dear_select)
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.internal_title:
                    subject = self.to_thai_number(str(doc_ids.internal_title))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                    if for_document == 'other':
                        for_document_other = doc_ids.for_document_other
                if doc_ids.material:
                    material = doc_ids.material
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if obj.job_title:
                    job_title = str(obj.job_title)
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                    day = self.to_thai_number(str(int((doc_ids.date_document).strftime('%d'))))
                    month = (int((doc_ids.date_document).strftime('%m')))
                    year = self.to_thai_number(str(int((doc_ids.date_document).strftime('%Y')) + 543))

                    if doc_ids.date_document_real:
                        date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document_real)))
                        day = self.to_thai_number(str(int((doc_ids.date_document_real).strftime('%d'))))
                        month = (int((doc_ids.date_document_real).strftime('%m')))
                        year = self.to_thai_number(str(int((doc_ids.date_document_real).strftime('%Y')) + 543))

                    if month == 1:
                        month_th = 'มกราคม'
                    if month == 2:
                        month_th = 'กุมภาพันธ์'
                    if month == 3:
                        month_th = 'มีนาคม'
                    if month == 4:
                        month_th = 'เมษายน'
                    if month == 5:
                        month_th = 'พฤษภาคม'
                    if month == 6:
                        month_th = 'มิถุนายน'
                    if month == 7:
                        month_th = 'กรกฎาคม'
                    if month == 8:
                        month_th = 'สิงหาคม'
                    if month == 9:
                        month_th = 'กันยายน'
                    if month == 10:
                        month_th = 'ตุลาคม'
                    if month == 11:
                        month_th = 'พฤศจิกายน'
                    if month == 12:
                        month_th = 'ธันวาคม'

                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                if doc_ids.date_document_real:
                    real_date = str(self.year_convert2(doc_ids.date_document_real))
                for employee_name in doc_ids.empolyee_name:
                    employee.append(str(employee_name.name))
                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)
                        employee_line = str(doc_line.employee_id.name)
                        employee_obj = self.env['hr.employee'].search([
                            ('name', '=', employee_line),
                        ])
                        employee_job = str(employee_obj.job_id.name)
                for line_dear in doc_ids.dear_select:
                    dear_name.append(str(line_dear))
                    receipt_group = []
                for receipt in doc_ids.reference_receive_document_related:
                    receipt_group.append(receipt)
                attached_file_in = []
                for attached in doc_ids.attached_file_in:
                    attached_file_in.append(attached)

                user_create = self.to_thai_number(str(doc_ids.create_uid.name))
                user_create_obj = self.env['hr.employee'].search([
                    ('user_id', '=', doc_ids.create_uid.id),
                ])

                department_create = self.to_thai_number(str(user_create_obj.department_id.name))
                department_email = self.to_thai_number(str(user_create_obj.work_email))

                department_obj = self.env['hr.department'].search([
                    ('name', '=', doc_ids.department_name),
                ])
                for department_line in department_obj:
                    if department_line.depa_phone:
                        department_phone = self.to_thai_number(str(department_line.depa_phone))
                    if department_line.depa_fax:
                        department_fax = self.to_thai_number(str(department_line.depa_fax))
                    # if department_line.depa_email:
                    #     department_email = self.to_thai_number(str(department_line.depa_email))

                # document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                # last_paragraph = document.paragraphs[-1]
                # last_paragraph.alignment = WD_TABLE_ALIGNMENT.RIGHT

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                # edit_record3 = document.add_paragraph('')
                # edit_record3.add_run('')
                # edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # edit_record3.paragraph_format.space_after = Pt(0)
                # edit_record4 = document.add_paragraph('')
                # edit_record4.add_run('')
                # edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # edit_record4.paragraph_format.space_after = Pt(0)

                zero_record = document.add_paragraph('')
                zero_record.style = document.styles['depa_speed']
                font51 = zero_record.add_run()
                font51.add_text(speed)
                zero_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                zero_record.paragraph_format.space_after = Pt(0)

                one_record = document.add_paragraph('ที่   ')
                one_record.add_run(real_name)
                one_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                one_record.paragraph_format.space_after = Pt(0)

                one_records = document.add_paragraph('')
                one_records.add_run('')
                one_records.add_run(day)
                one_records.add_run(' ')
                one_records.add_run(month_th)
                one_records.add_run(' ')
                one_records.add_run(year)
                one_records.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                one_records.paragraph_format.space_after = Pt(0)
                one_records.paragraph_format.space_before = Pt(6)
                one_records.paragraph_format.left_indent = Inches(3.149606)

                two_record = document.add_paragraph()
                two_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                two_record.add_run('เรื่อง').bold = True
                two_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                two_record.paragraph_format.space_after = Pt(0)
                two_record.paragraph_format.space_before = Pt(6)
                two_record.add_run("\t")
                two_record.add_run(subject)
                two_record.paragraph_format.left_indent = Inches(0.4)
                two_record.paragraph_format.first_line_indent = Inches(-0.4)


                three_record = document.add_paragraph()
                three_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                three_record.add_run('เรียน').bold = True
                three_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                three_record.paragraph_format.space_after = Pt(0)
                three_record.paragraph_format.space_before = Pt(6)
                three_record.add_run("\t")
                three_record.add_run(dear_name)
                three_record.paragraph_format.left_indent = Inches(0.4)
                three_record.paragraph_format.first_line_indent = Inches(-0.4)

                if len(receipt_group) > 0:
                    four_record = document.add_paragraph('อ้างถึง')
                    four_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    four_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    four_record.paragraph_format.space_after = Pt(0)
                    four_record.paragraph_format.space_before = Pt(0)
                    four_record.add_run("\t")
                    four_record.add_run('หนังสือ')
                    # if receipt_group[0].from_document == False:
                    #     four_record.add_run(' ')
                    if receipt_group[0].from_document != False:
                        four_record.add_run(' '+receipt_group[0].from_document)
                    if receipt_group[0].refer != False:
                        four_record.add_run(' ที่ ')
                        four_record.add_run(self.to_thai_number(str(receipt_group[0].refer)))
                    four_record.add_run(' ลงวันที่ ')
                    four_record.add_run(
                        self.to_thai_number(str(self.to_full_thai_date(receipt_group[0].date_receive))))
                    four_record.paragraph_format.left_indent = Inches(0.5)
                    four_record.paragraph_format.first_line_indent = Inches(-0.5)
                    if len(receipt_group) > 1:
                        four_record1 = document.add_paragraph("\t")
                        four_record1.add_run('หนังสือ')
                        # if receipt_group[1].from_document == False:
                        #     four_record1.add_run(' ')
                        if receipt_group[1].from_document != False:
                            four_record1.add_run(' '+receipt_group[1].from_document)
                        if receipt_group[1].refer != False:
                            four_record1.add_run(' ที่ ')
                            four_record1.add_run(self.to_thai_number(str(receipt_group[1].refer)))
                        four_record1.add_run(' ลงวันที่ ')
                        four_record1.add_run(
                            self.to_thai_number(str(self.to_full_thai_date(receipt_group[1].date_receive))))
                        four_record1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        four_record1.paragraph_format.space_after = Pt(0)
                        four_record1.paragraph_format.space_before = Pt(0)
                        four_record1.paragraph_format.left_indent = Inches(0.5)
                        four_record1.paragraph_format.first_line_indent = Inches(-0.5)
                        if len(receipt_group) > 2:
                            four_record2 = document.add_paragraph("\t")
                            four_record2.add_run('หนังสือ')
                            # if receipt_group[2].from_document == False:
                            #     four_record2.add_run(' ')
                            if receipt_group[2].from_document != False:
                                four_record2.add_run(' '+receipt_group[2].from_document)
                            if receipt_group[2].refer != False:
                                four_record2.add_run(' ที่ ')
                                four_record2.add_run(self.to_thai_number(str(receipt_group[2].refer)))
                            four_record2.add_run(' ลงวันที่ ')
                            four_record2.add_run(
                                self.to_thai_number(str(self.to_full_thai_date(receipt_group[2].date_receive))))
                            four_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            four_record2.paragraph_format.space_after = Pt(0)
                            four_record2.paragraph_format.space_before = Pt(0)
                            four_record2.paragraph_format.left_indent = Inches(0.5)
                            four_record2.paragraph_format.first_line_indent = Inches(-0.5)
                            if len(receipt_group) > 3:
                                four_record3 = document.add_paragraph("\t")
                                four_record3.add_run('หนังสือ')
                                # if receipt_group[3].from_document == False:
                                #     four_record3.add_run(' ')
                                if receipt_group[3].from_document != False:
                                    four_record3.add_run(' '+receipt_group[3].from_document)
                                if receipt_group[3].refer != False:
                                    four_record3.add_run(' ที่ ')
                                    four_record3.add_run(self.to_thai_number(str(receipt_group[3].refer)))
                                four_record3.add_run(' ลงวันที่ ')
                                four_record3.add_run(
                                    self.to_thai_number(str(self.to_full_thai_date(receipt_group[3].date_receive))))
                                four_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                                four_record3.paragraph_format.space_after = Pt(0)
                                four_record3.paragraph_format.space_before = Pt(0)
                                four_record3.paragraph_format.left_indent = Inches(0.5)
                                four_record3.paragraph_format.first_line_indent = Inches(-0.5)

                if len(attached_file_in) > 0:
                    seq = ''
                    if len(attached_file_in) > 1:
                        seq = self.to_single_thai_number(1) + '. '
                    five_record = document.add_paragraph(f'สิ่งที่ส่งมาด้วย')
                    five_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    five_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    five_record.paragraph_format.space_after = Pt(0)
                    five_record.paragraph_format.space_before = Pt(6)
                    five_record.add_run("\t")
                    five_record.add_run(f'{str(seq)}{str(attached_file_in[0].name)}')
                    five_record.paragraph_format.left_indent = Inches(0.9)
                    five_record.paragraph_format.first_line_indent = Inches(-0.9)
                    if len(attached_file_in) > 1:
                        for i in range(1, len(attached_file_in)):
                            five_record1 = document.add_paragraph("\t")
                            five_record1.add_run(f'{self.to_single_thai_number(i+1)}. {str(attached_file_in[i].name)}')
                            five_record1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            five_record1.paragraph_format.space_after = Pt(0)
                            five_record1.paragraph_format.space_before = Pt(0)
                            five_record1.paragraph_format.left_indent = Inches(0.9)
                            five_record1.paragraph_format.first_line_indent = Inches(-0.9)
                        #     five_record1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        #     five_record1.paragraph_format.space_after = Pt(0)
                        #     five_record1.paragraph_format.space_before = Pt(6)
                        # if len(attached_file_in) > 2:
                        #     five_record2 = document.add_paragraph('                   ')
                        #     five_record2.add_run(self.to_thai_number(str(attached_file_in[2].name)))
                        #     five_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        #     five_record2.paragraph_format.space_after = Pt(0)
                        #     five_record2.paragraph_format.space_before = Pt(6)
                        #     if len(attached_file_in) > 3:
                        #         five_record3 = document.add_paragraph('                  ')
                        #         five_record3.add_run(self.to_thai_number(str(attached_file_in[3].name)))
                        #         five_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        #         five_record3.paragraph_format.space_after = Pt(0)
                        #         five_record3.paragraph_format.space_before = Pt(6)
                    # five_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    # five_record.paragraph_format.space_after = Pt(0)
                    # five_record.paragraph_format.space_before = Pt(0)


                if "\n" in material.strip():
                    for p in material.split("\n"):
                        if p:
                            six_record = document.add_paragraph('')
                            six_record.add_run(p)
                            six_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            six_record.paragraph_format.space_after = Pt(0)
                            six_record.paragraph_format.space_before = Pt(6)
                            six_record.paragraph_format.first_line_indent = Inches(0.9842520)
                else:
                    six_record = document.add_paragraph('')
                    six_record.add_run(material)
                    six_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    six_record.paragraph_format.space_after = Pt(0)
                    six_record.paragraph_format.space_before = Pt(6)
                    six_record.paragraph_format.first_line_indent = Inches(0.9842520)

                if for_document != 'other':
                    seven_record = document.add_paragraph('จึงเรียนมา')
                    seven_record.add_run(for_document)
                else:
                    seven_record = document.add_paragraph('')
                    seven_record.add_run(for_document_other)
                seven_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                seven_record.paragraph_format.space_after = Pt(0)
                seven_record.paragraph_format.space_before = Pt(6)
                seven_record.paragraph_format.left_indent = Inches(0.9842520)

                seven_record2 = document.add_paragraph('ขอแสดงความนับถือ')
                seven_record2.add_run('')
                seven_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                seven_record2.paragraph_format.space_after = Pt(0)
                seven_record2.paragraph_format.space_before = Pt(6)
                seven_record2.paragraph_format.left_indent = Inches(3.149606)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record3 = document.add_paragraph('')
                edit_record3.add_run('')
                edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record3.paragraph_format.space_after = Pt(0)
                edit_record4 = document.add_paragraph('')
                edit_record4.add_run('')
                edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record4.paragraph_format.space_after = Pt(0)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(employee_line)
                sevens_record.add_run(' )')
                sevens_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                sevens_record.paragraph_format.space_after = Pt(0)
                sevens_record.paragraph_format.space_before = Pt(6)
                sevens_record.paragraph_format.left_indent = Inches(3)
                # sevens_record.paragraph_format.left_indent = Inches(2.952756)

                eight_record = document.add_paragraph('')
                eight_record.add_run(employee_job)
                eight_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                eight_record.paragraph_format.space_after = Pt(0)
                eight_record.paragraph_format.space_before = Pt(6)
                eight_record.paragraph_format.left_indent = Inches(2.5)
                # eight_record.paragraph_format.left_indent = Inches(2.460630)

                edit_record5 = document.add_paragraph('')
                edit_record5.add_run('')
                edit_record5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record5.paragraph_format.space_after = Pt(0)
                # edit_record6 = document.add_paragraph('')
                # edit_record6.add_run('')
                # edit_record6.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # edit_record6.paragraph_format.space_after = Pt(0)
                # edit_record7 = document.add_paragraph('')
                # edit_record7.add_run('')
                # edit_record7.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # edit_record7.paragraph_format.space_after = Pt(0)

                nine_record = document.add_paragraph('')
                nine_record.add_run(department_create)
                nine_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                nine_record.paragraph_format.space_after = Pt(0)
                nine_record.paragraph_format.space_before = Pt(0)

                ten_record = document.add_paragraph('โทรศัพท์  ')
                # ten_record.add_run(department_phone)
                ten_record.add_run(self.to_thai_number('02-026-2333'))
                ten_record.add_run('   (')
                ten_record.add_run(user_create)
                ten_record.add_run(')')
                ten_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                ten_record.paragraph_format.space_after = Pt(0)
                ten_record.paragraph_format.space_before = Pt(0)

                twelve_record = document.add_paragraph('ไปรษณีย์อิเล็กทรอนิกส์ : ')
                twelve_record.add_run(department_email)
                twelve_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                twelve_record.paragraph_format.space_after = Pt(0)
                twelve_record.paragraph_format.space_before = Pt(0)

                # font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run(name).font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run('  ').font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run(user_name).font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run('  ').font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run(user_job_title).font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run('  ').font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run(current_date).font
                # font.size = Pt(10)
            if doc_ids.document_type == 'ระเบียบ':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.subject:
                    subject = self.to_thai_number(str(doc_ids.subject))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                if doc_ids.material:
                    material = str(doc_ids.material)
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                    day = self.to_thai_number(str(int((doc_ids.date_document).strftime('%d'))))
                    month = (int((doc_ids.date_document).strftime('%m')))
                    if month == 1:
                        month_th = 'มกราคม'
                    if month == 2:
                        month_th = 'กุมภาพันธ์'
                    if month == 3:
                        month_th = 'มีนาคม'
                    if month == 4:
                        month_th = 'เมษายน'
                    if month == 5:
                        month_th = 'พฤษภาคม'
                    if month == 6:
                        month_th = 'มิถุนายน'
                    if month == 7:
                        month_th = 'กรกฎาคม'
                    if month == 8:
                        month_th = 'สิงหาคม'
                    if month == 9:
                        month_th = 'กันยายน'
                    if month == 10:
                        month_th = 'ตุลาคม'
                    if month == 11:
                        month_th = 'พฤศจิกายน'
                    if month == 12:
                        month_th = 'ธันวาคม'
                    year = self.to_thai_number(str(int((doc_ids.date_document).strftime('%Y')) + 543))
                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                    len_real_name = str(len(doc_ids.name_real))
                    print(len_real_name)
                for employee_name in doc_ids.empolyee_name:
                    employee.append(str(employee_name.name))
                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)
                        employee_line = str(doc_line.employee_id.name)
                        employee_job = str(doc_line.job_id_name.name)
                department_obj = self.env['hr.department'].search([
                    ('name', '=', doc_ids.department_name),
                ])
                for department_line in department_obj:
                    if department_line.depa_phone:
                        department_phone = department_line.depa_phone
                    if department_line.depa_fax:
                        department_fax = department_line.depa_fax
                    if department_line.depa_email:
                        department_email = department_line.depa_email
                sing_department = self.to_thai_number(str(doc_ids.sign_department.name))
                sing_employee = self.to_thai_number(str(doc_ids.sign_employee.name))

                document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record = document.add_paragraph('')
                font = one_record.add_run('ระเบียบสำนักงานส่งเสริมเศรษฐกิจดิจิทัล').font
                font.size = Pt(18)
                font.bold = True
                one_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record.paragraph_format.space_after = Pt(0)
                one_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                two_record = document.add_paragraph('')
                font1 = two_record.add_run('ว่าด้วย ').font
                font2 = two_record.add_run(subject).font
                font3 = two_record.add_run(' ').font
                font4 = two_record.add_run(date).font
                font1.size = Pt(18)
                font1.bold = True
                font2.size = Pt(18)
                font2.bold = True
                font3.size = Pt(18)
                font3.bold = True
                font4.size = Pt(18)
                font4.bold = True
                two_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                two_record.paragraph_format.space_after = Pt(0)
                two_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                three_record = document.add_paragraph('')
                three_record.add_run('ฉบับที่ ').bold = True
                if len_real_name == '14':
                    font1 = three_record.add_run(real_name[8]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '15':
                    font1 = three_record.add_run(real_name[8:10]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '16':
                    font1 = three_record.add_run(real_name[8:11]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '17':
                    font1 = three_record.add_run(real_name[8:12]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '18':
                    font1 = three_record.add_run(real_name[8:13]).font
                    font1.size = Pt(18)
                    font1.bold = True
                three_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                three_record.paragraph_format.space_after = Pt(0)
                three_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                four_record = document.add_paragraph('')
                four_record.add_run('พ.ศ. ').bold = True
                if len_real_name == '14':
                    font1 = four_record.add_run(real_name[10:14]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '15':
                    font1 = four_record.add_run(real_name[11:15]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '16':
                    font1 = four_record.add_run(real_name[12:16]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '17':
                    font1 = four_record.add_run(real_name[13:17]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '18':
                    font1 = four_record.add_run(real_name[14:18]).font
                    font1.size = Pt(18)
                    font1.bold = True

                four_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                four_record.paragraph_format.space_after = Pt(0)
                four_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                five_record = document.add_paragraph('____________________________________  ')
                five_record.add_run('')
                five_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                edit_record = document.add_paragraph('')
                edit_record.add_run('')
                edit_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record.paragraph_format.space_after = Pt(0)
                edit_record.paragraph_format.space_before = Pt(12)

                six_record = document.add_paragraph('')
                six_record.add_run(material)
                six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                six_record.paragraph_format.space_after = Pt(0)
                six_record.paragraph_format.space_before = Pt(6)

                # edit_record2 = document.add_paragraph('')
                # edit_record2.add_run('')
                # edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # edit_record2.paragraph_format.space_after = Pt(0)
                # edit_record2.paragraph_format.space_before = Pt(12)

                seven_record = document.add_paragraph('ประกาศ ณ วันที่ ')
                seven_record.add_run(day)
                seven_record.add_run(' ')
                seven_record.add_run('เดือน ')
                seven_record.add_run(month_th)
                seven_record.add_run(' ')
                seven_record.add_run('พ.ศ. ')
                seven_record.add_run(year)
                seven_record.paragraph_format.left_indent = Inches(0.9842520)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record3 = document.add_paragraph('')
                edit_record3.add_run('')
                edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record3.paragraph_format.space_after = Pt(0)
                edit_record4 = document.add_paragraph('')
                edit_record4.add_run('')
                edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record4.paragraph_format.space_after = Pt(0)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(sing_employee)
                sevens_record.add_run(' )')
                sevens_record.paragraph_format.left_indent = Inches(3.149606)

                eight_record = document.add_paragraph('')
                eight_record.add_run(sing_department)
                eight_record.paragraph_format.left_indent = Inches(3.051181)
                edits_record = document.add_paragraph('')
                edits_record.add_run('')

                font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_job_title).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(current_date).font
                font.size = Pt(10)
            if doc_ids.document_type == 'ประกาศ':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                sing_department = self.to_thai_number(str(doc_ids.sign_department.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                speed = self.text_speed(str(doc_ids.speed))
                secret = self.text_secret(str(doc_ids.secret))
                print(secret, speed)
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.subject:
                    subject = self.to_thai_number(str(doc_ids.subject))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                if doc_ids.material:
                    material = str(doc_ids.material)
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                    day = self.to_thai_number(str(int((doc_ids.date_document).strftime('%d'))))
                    month = (int((doc_ids.date_document).strftime('%m')))
                    if month == 1:
                        month_th = 'มกราคม'
                    if month == 2:
                        month_th = 'กุมภาพันธ์'
                    if month == 3:
                        month_th = 'มีนาคม'
                    if month == 4:
                        month_th = 'เมษายน'
                    if month == 5:
                        month_th = 'พฤษภาคม'
                    if month == 6:
                        month_th = 'มิถุนายน'
                    if month == 7:
                        month_th = 'กรกฎาคม'
                    if month == 8:
                        month_th = 'สิงหาคม'
                    if month == 9:
                        month_th = 'กันยายน'
                    if month == 10:
                        month_th = 'ตุลาคม'
                    if month == 11:
                        month_th = 'พฤศจิกายน'
                    if month == 12:
                        month_th = 'ธันวาคม'
                    year = self.to_thai_number(str(int((doc_ids.date_document).strftime('%Y')) + 543))
                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                    len_real_name = str(len(doc_ids.name_real))
                for employee_name in doc_ids.empolyee_name:
                    employee.append(str(employee_name.name))
                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)
                        employee_line = str(doc_line.employee_id.name)
                        employee_job = str(doc_line.job_id_name.name)
                department_obj = self.env['hr.department'].search([
                    ('name', '=', doc_ids.department_name),
                ])
                for department_line in department_obj:
                    if department_line.depa_phone:
                        department_phone = department_line.depa_phone
                    if department_line.depa_fax:
                        department_fax = department_line.depa_fax
                    if department_line.depa_email:
                        department_email = department_line.depa_email

                document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                zero_record = document.add_paragraph('')
                zero_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                zero_record.paragraph_format.space_after = Pt(0)
                font1 = zero_record.add_run('ประกาศ').font
                font1.size = Pt(18)
                font1.bold = True
                zero_record.add_run('')
                zero_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                one_record = document.add_paragraph('')
                font1 = one_record.add_run('สำนักงานส่งเสริมเศรษฐกิจดิจิทัล').font
                font1.size = Pt(18)
                font1.bold = True
                one_record.add_run('')
                one_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                one_record.paragraph_format.space_after = Pt(0)
                one_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                two_record = document.add_paragraph('')
                font = two_record.add_run('เรื่อง ').font
                font2 = two_record.add_run(subject).font
                font3 = two_record.add_run(' ').font
                font4 = two_record.add_run(date).font
                font.size = Pt(18)
                font.bold = True
                font2.size = Pt(18)
                font2.bold = True
                font3.size = Pt(18)
                font3.bold = True
                font4.size = Pt(18)
                font4.bold = True
                two_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                two_record.paragraph_format.space_after = Pt(0)
                two_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                five_record = document.add_paragraph('____________________________________  ')
                five_record.add_run('')
                five_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                edit_record = document.add_paragraph('')
                edit_record.add_run('')
                edit_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record.paragraph_format.space_after = Pt(0)
                edit_record.paragraph_format.space_before = Pt(12)

                six_record = document.add_paragraph('')
                six_record.add_run(material)
                six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                six_record.paragraph_format.space_after = Pt(6)
                six_record.paragraph_format.space_before = Pt(0)

                seven_record = document.add_paragraph('ประกาศ ณ วันที่ ')
                seven_record.add_run(day)
                seven_record.add_run(' ')
                seven_record.add_run('เดือน ')
                seven_record.add_run(month_th)
                seven_record.add_run(' ')
                seven_record.add_run('พ.ศ. ')
                seven_record.add_run(year)
                seven_record.paragraph_format.left_indent = Inches(0.9842520)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record3 = document.add_paragraph('')
                edit_record3.add_run('')
                edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record3.paragraph_format.space_after = Pt(0)
                edit_record4 = document.add_paragraph('')
                edit_record4.add_run('')
                edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record4.paragraph_format.space_after = Pt(0)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(employee_line)
                sevens_record.add_run(' )')
                sevens_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                eight_record = document.add_paragraph('')
                eight_record.add_run(sing_department)
                eight_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                edits_record = document.add_paragraph('')
                edits_record.add_run('')
                nine_record = document.add_paragraph('')
                nine_record.add_run(note)

                font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_job_title).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(current_date).font
                font.size = Pt(10)
            if doc_ids.document_type == 'ข้อบังคับ':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                speed = self.text_speed(str(doc_ids.speed))
                secret = self.text_secret(str(doc_ids.secret))
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.subject:
                    subject = self.to_thai_number(str(doc_ids.subject))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                if doc_ids.material:
                    material = str(doc_ids.material)
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                    day = self.to_thai_number(str(int((doc_ids.date_document).strftime('%d'))))
                    month = (int((doc_ids.date_document).strftime('%m')))
                    if month == 1:
                        month_th = 'มกราคม'
                    if month == 2:
                        month_th = 'กุมภาพันธ์'
                    if month == 3:
                        month_th = 'มีนาคม'
                    if month == 4:
                        month_th = 'เมษายน'
                    if month == 5:
                        month_th = 'พฤษภาคม'
                    if month == 6:
                        month_th = 'มิถุนายน'
                    if month == 7:
                        month_th = 'กรกฎาคม'
                    if month == 8:
                        month_th = 'สิงหาคม'
                    if month == 9:
                        month_th = 'กันยายน'
                    if month == 10:
                        month_th = 'ตุลาคม'
                    if month == 11:
                        month_th = 'พฤศจิกายน'
                    if month == 12:
                        month_th = 'ธันวาคม'
                    year = self.to_thai_number(str(int((doc_ids.date_document).strftime('%Y')) + 543))
                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                    len_real_name = str(len(doc_ids.name_real))
                for employee_name in doc_ids.empolyee_name:
                    employee.append(str(employee_name.name))
                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)
                        employee_line = str(doc_line.employee_id.name)
                        employee_job = str(doc_line.job_id_name.name)
                department_obj = self.env['hr.department'].search([
                    ('name', '=', doc_ids.department_name),
                ])
                for department_line in department_obj:
                    if department_line.depa_phone:
                        department_phone = department_line.depa_phone
                    if department_line.depa_fax:
                        department_fax = department_line.depa_fax
                    if department_line.depa_email:
                        department_email = department_line.depa_email

                sing_department = self.to_thai_number(str(doc_ids.sign_department.name))
                sing_employee = self.to_thai_number(str(doc_ids.sign_employee.name))

                document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record = document.add_paragraph('')
                font = one_record.add_run('ข้อบังคับสำนักงานส่งเสริมเศรษฐกิจดิจิทัล').font
                font.size = Pt(18)
                font.bold = True
                one_record.add_run('')
                one_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                one_record.paragraph_format.space_after = Pt(0)

                two_record = document.add_paragraph('')
                font = two_record.add_run('ว่าด้วย ').font
                font2 = two_record.add_run(subject).font
                font3 = two_record.add_run(' ').font
                font4 = two_record.add_run(date).font
                font.size = Pt(18)
                font.bold = True
                font2.size = Pt(18)
                font2.bold = True
                font3.size = Pt(18)
                font3.bold = True
                font4.size = Pt(18)
                font4.bold = True
                two_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                two_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                two_record.paragraph_format.space_after = Pt(0)

                three_record = document.add_paragraph('')
                three_record.add_run('ฉบับที่ ').bold = True
                if len_real_name == '16':
                    font1 = three_record.add_run(real_name[10]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '17':
                    font1 = three_record.add_run(real_name[10:12]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '18':
                    font1 = three_record.add_run(real_name[10:13]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '19':
                    font1 = three_record.add_run(real_name[10:14]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '20':
                    font1 = three_record.add_run(real_name[10:15]).font
                    font1.size = Pt(18)
                    font1.bold = True
                three_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                three_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                three_record.paragraph_format.space_after = Pt(0)

                four_record = document.add_paragraph('')
                four_record.add_run('พ.ศ. ').bold = True
                if len_real_name == '14':
                    font1 = four_record.add_run(real_name[10:14]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '15':
                    font1 = four_record.add_run(real_name[11:15]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '16':
                    font1 = four_record.add_run(real_name[12:16]).font
                    font1.size = Pt(18)
                    font1.bold = True

                if len_real_name == '17':
                    font1 = four_record.add_run(real_name[13:17]).font
                    font1.size = Pt(18)
                    font1.bold = True

                if len_real_name == '18':
                    font1 = four_record.add_run(real_name[14:18]).font
                    font1.size = Pt(18)
                    font1.bold = True

                four_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                four_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                four_record.paragraph_format.space_after = Pt(0)

                five_record = document.add_paragraph('____________________________________  ')
                five_record.add_run('')
                five_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                edit_record = document.add_paragraph('')
                edit_record.add_run('')
                edit_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record.paragraph_format.space_after = Pt(0)
                edit_record.paragraph_format.space_before = Pt(12)

                six_record = document.add_paragraph('')
                six_record.add_run(material)
                six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                six_record.paragraph_format.space_after = Pt(6)
                six_record.paragraph_format.space_before = Pt(0)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record2.paragraph_format.space_before = Pt(12)

                seven_record = document.add_paragraph('ประกาศ ณ วันที่ ')
                seven_record.add_run(day)
                seven_record.add_run(' ')
                seven_record.add_run('เดือน ')
                seven_record.add_run(month_th)
                seven_record.add_run(' ')
                seven_record.add_run('พ.ศ. ')
                seven_record.add_run(year)
                seven_record.paragraph_format.left_indent = Inches(0.9842520)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record3 = document.add_paragraph('')
                edit_record3.add_run('')
                edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record3.paragraph_format.space_after = Pt(0)
                edit_record4 = document.add_paragraph('')
                edit_record4.add_run('')
                edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record4.paragraph_format.space_after = Pt(0)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(sing_employee)
                sevens_record.add_run(' )')
                sevens_record.paragraph_format.left_indent = Inches(3.149606)

                eight_record = document.add_paragraph('')
                eight_record.add_run(sing_department)
                eight_record.paragraph_format.left_indent = Inches(3.051181)
                edits_record = document.add_paragraph('')
                edits_record.add_run('')

                font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_job_title).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(current_date).font
                font.size = Pt(10)

        fp = BytesIO()
        document.save(fp)
        fp.seek(0)
        data = fp.read()
        fp.close()
        return data

    def from_data2(self, current_id):
        cur_path = os.path.dirname(os.path.abspath(__file__))
        docx_path = cur_path + '/template/format_document.docx'
        rec_ids = self.env['document.internal.main'].search([('id', 'in', [current_id])])
        document = Document()
        # document = DocxTemplate(docx_path)
        style = document.styles['Normal']
        font = style.font
        font.name = 'TH SarabunPSK'
        font.size = Pt(16)
        section = document.sections[0]
        footer = section.footer
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Inches(1.181102)
        section.right_margin = Inches(0.7874016)
        section.top_margin = Inches(0.3937007)
        section.bottom_margin = Inches(0.3937007)
        # section.top_margin = Inches(0.9842520)
        # section.bottom_margin = Inches(0.5905512)
        date = ''
        real_date = ''
        depa_department = ''
        depa_phone = ''
        subject = ''
        for_document = ''
        for_document_other = ''
        material = ''
        note = ''
        job_title = ''
        real_name = ''
        department_phone = ''
        department_fax = ''
        department_email = ''
        employee = []
        state_line2 = []
        dear_name = []
        # styles = document.styles
        # paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
        # for style in paragraph_styles:
        #     print(style.name)

        for doc_ids in rec_ids:
            if doc_ids.document_type == 'คำสั่ง ก' or doc_ids.document_type == 'คำสั่ง ข' or \
                    doc_ids.document_type == 'คำสั่ง ค' or doc_ids.document_type == 'คำสั่ง พ':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                speed = self.text_speed(str(doc_ids.speed))
                secret = self.text_secret(str(doc_ids.secret))
                print(secret, speed)
                sing_department = self.to_thai_number(str(doc_ids.sign_department.name))
                sing_employee = self.to_thai_number(str(doc_ids.sign_employee.name))
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.subject:
                    subject = self.to_thai_number(str(doc_ids.subject))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                if doc_ids.material:
                    material = str(doc_ids.material)
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                    day = self.to_thai_number(str(int((doc_ids.date_document).strftime('%d'))))
                    month = (int((doc_ids.date_document).strftime('%m')))
                    if month == 1:
                        month_th = 'มกราคม'
                    if month == 2:
                        month_th = 'กุมภาพันธ์'
                    if month == 3:
                        month_th = 'มีนาคม'
                    if month == 4:
                        month_th = 'เมษายน'
                    if month == 5:
                        month_th = 'พฤษภาคม'
                    if month == 6:
                        month_th = 'มิถุนายน'
                    if month == 7:
                        month_th = 'กรกฎาคม'
                    if month == 8:
                        month_th = 'สิงหาคม'
                    if month == 9:
                        month_th = 'กันยายน'
                    if month == 10:
                        month_th = 'ตุลาคม'
                    if month == 11:
                        month_th = 'พฤศจิกายน'
                    if month == 12:
                        month_th = 'ธันวาคม'
                    year = self.to_thai_number(str(int((doc_ids.date_document).strftime('%Y')) + 543))
                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                    len_real_name = str(len(doc_ids.name_real))
                for employee_name in doc_ids.empolyee_name:
                    employee.append(str(employee_name.name))
                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)
                        employee_line = str(doc_line.employee_id.name)
                        employee_job = str(doc_line.job_id_name.name)
                department_obj = self.env['hr.department'].search([
                    ('name', '=', doc_ids.department_name),
                ])
                for department_line in department_obj:
                    if department_line.depa_phone:
                        department_phone = department_line.depa_phone
                    if department_line.depa_fax:
                        department_fax = department_line.depa_fax
                    if department_line.depa_email:
                        department_email = department_line.depa_email

                document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER

                zero_record = document.add_paragraph('')
                # font1 = zero_record.add_run('คำสั่ง').font
                # font1.size = Pt(18)
                # font1.bold = True
                zero_record.add_run('')
                zero_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                zero_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                zero_record.paragraph_format.space_after = Pt(0)

                one_record = document.add_paragraph('')
                font1 = one_record.add_run('คำสั่งสำนักงานส่งเสริมเศรษฐกิจดิจิทัล').font
                font1.size = Pt(18)
                font1.bold = True
                one_record.add_run('')
                one_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                one_record.paragraph_format.space_after = Pt(0)

                two_record = document.add_paragraph('')
                font1 = two_record.add_run('ที่  ').font
                font2 = two_record.add_run(name).font
                font1.size = Pt(18)
                font1.bold = True
                font2.size = Pt(18)
                font2.bold = True
                two_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                two_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                two_record.paragraph_format.space_after = Pt(0)

                three_record = document.add_paragraph('')
                font1 = three_record.add_run('เรื่อง ')
                font2 = three_record.add_run(subject).font
                font1.size = Pt(18)
                font1.bold = True
                font2.size = Pt(18)
                font2.bold = True
                three_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                three_record.paragraph_format.space_after = Pt(0)
                three_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                five_record = document.add_paragraph('____________________________________  ')
                five_record.add_run('')
                five_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                edit_record = document.add_paragraph('')
                edit_record.add_run('')
                edit_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record.paragraph_format.space_after = Pt(0)
                edit_record.paragraph_format.space_before = Pt(12)

                six_record = document.add_paragraph('')
                six_record.add_run(material)
                six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                six_record.paragraph_format.space_after = Pt(0)
                six_record.paragraph_format.space_before = Pt(6)

                six_record2 = document.add_paragraph('ทั้งนี้ ตั้งแต่บัดนี้เป็นต้นไป')
                six_record2.add_run('')
                six_record2.paragraph_format.left_indent = Inches(0.9842520)
                six_record2.paragraph_format.space_after = Pt(0)
                six_record2.paragraph_format.space_before = Pt(6)

                seven_record = document.add_paragraph('สั่ง ณ วันที่ ')
                seven_record.add_run(day)
                seven_record.add_run(' ')
                seven_record.add_run('เดือน ')
                seven_record.add_run(month_th)
                seven_record.add_run(' ')
                seven_record.add_run('พ.ศ. ')
                seven_record.add_run(year)
                seven_record.paragraph_format.left_indent = Inches(1.574803)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record3 = document.add_paragraph('')
                edit_record3.add_run('')
                edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record3.paragraph_format.space_after = Pt(0)
                edit_record4 = document.add_paragraph('')
                edit_record4.add_run('')
                edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record4.paragraph_format.space_after = Pt(0)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(sing_employee)
                sevens_record.add_run(' )')
                sevens_record.paragraph_format.left_indent = Inches(3.149606)

                eight_record = document.add_paragraph('')
                eight_record.add_run(sing_department)
                eight_record.paragraph_format.left_indent = Inches(3.051181)
                edits_record = document.add_paragraph('')
                edits_record.add_run('')

                font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_job_title).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(current_date).font
                font.size = Pt(10)
            if doc_ids.document_type == 'บันทึกข้อความ':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.subject:
                    subject = self.to_thai_number(str(doc_ids.subject))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                if doc_ids.material:
                    material = str(doc_ids.material)
                state = str(doc_ids.state)
                sign = str(doc_ids.sign.name)
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if obj.job_title:
                    job_title = str(obj.job_title)
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                if doc_ids.date_document_real:
                    real_date = str(self.year_convert2(doc_ids.date_document_real))
                job_em = []
                for employee_name in doc_ids.setting_line_ids_related:
                    if employee_name.is_active:
                        employee.append(str(employee_name.employee_id.name))
                        job_em.append(str(employee_name.employee_id.job_title))
                print(employee_name, employee, job_em)

                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)

                document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                header_fist = document.add_paragraph('                                                 ').add_run(
                    'บันทึกข้อความ')
                header_fist.bold = True
                font = header_fist.font
                font.name = 'TH SarabunPSK'
                font.size = Pt(24)
                header_fist.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record = document.add_paragraph('ฝ่าย/ส่วน ')
                # paragraph_format = header_fist.paragraph_format
                # paragraph_format.line_spacing = Pt(18)
                one_record.underline = WD_UNDERLINE.DOUBLE
                one_record.add_run(depa_department)
                one_record.add_run('                                                             ')
                one_record.add_run('โทร. ')
                one_record.add_run('')
                one_record.add_run(depa_phone)
                two_record = document.add_paragraph('ที่ ')
                two_record.add_run(real_name)
                two_record.add_run(
                    '                                                   ')
                two_record.add_run('วันที่ ')
                two_record.add_run(date)
                three_record = document.add_paragraph('เรื่อง ')
                three_record.add_run(subject)
                four_record = document.add_paragraph('เรียน ')
                four_record.add_run(employee[-1])
                five_record = document.add_paragraph('ผ่าน ')
                for line_name2 in state_line2:
                    five_record.add_run(' ')
                    five_record.add_run(line_name2)
                    five_record.add_run(',')
                six_record = document.add_paragraph('')
                six_record.add_run(material)
                seven_record = document.add_paragraph('')
                seven_record = document.add_paragraph('จึงมาเรียน')
                seven_record.add_run(for_document)
                seven_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                seven_record.paragraph_format.left_indent = Inches(0.9842520)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(employee[-1])
                sevens_record.add_run(' )')
                sevens_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                sevens_record.paragraph_format.left_indent = Inches(3.444882)

                eight_record = document.add_paragraph('')
                eight_record.add_run(job_em[-1])
                eight_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                eight_record.paragraph_format.left_indent = Inches(3.346457)

                font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_job_title).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(current_date).font
                font.size = Pt(10)
            if doc_ids.document_type == 'หนังสือภายนอก+หนังสือรับรอง':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                speed = self.text_speed(str(doc_ids.speed))
                secret = self.text_secret(str(doc_ids.secret))
                # print(doc_ids.secret)
                dear_select = str(doc_ids.dear_select)
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.internal_title:
                    subject = self.to_thai_number(str(doc_ids.internal_title))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                    if for_document == 'other':
                        for_document_other = doc_ids.for_document_other
                if doc_ids.material:
                    material = doc_ids.material
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if obj.job_title:
                    job_title = str(obj.job_title)
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                    day = self.to_thai_number(str(int((doc_ids.date_document).strftime('%d'))))
                    month = (int((doc_ids.date_document).strftime('%m')))
                    year = self.to_thai_number(str(int((doc_ids.date_document).strftime('%Y')) + 543))

                    if doc_ids.date_document_real:
                        date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document_real)))
                        day = self.to_thai_number(str(int((doc_ids.date_document_real).strftime('%d'))))
                        month = (int((doc_ids.date_document_real).strftime('%m')))
                        year = self.to_thai_number(str(int((doc_ids.date_document_real).strftime('%Y')) + 543))

                    if month == 1:
                        month_th = 'มกราคม'
                    if month == 2:
                        month_th = 'กุมภาพันธ์'
                    if month == 3:
                        month_th = 'มีนาคม'
                    if month == 4:
                        month_th = 'เมษายน'
                    if month == 5:
                        month_th = 'พฤษภาคม'
                    if month == 6:
                        month_th = 'มิถุนายน'
                    if month == 7:
                        month_th = 'กรกฎาคม'
                    if month == 8:
                        month_th = 'สิงหาคม'
                    if month == 9:
                        month_th = 'กันยายน'
                    if month == 10:
                        month_th = 'ตุลาคม'
                    if month == 11:
                        month_th = 'พฤศจิกายน'
                    if month == 12:
                        month_th = 'ธันวาคม'

                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                if doc_ids.date_document_real:
                    real_date = str(self.year_convert2(doc_ids.date_document_real))
                for employee_name in doc_ids.empolyee_name:
                    employee.append(str(employee_name.name))
                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)
                        employee_line = str(doc_line.employee_id.name)
                        employee_obj = self.env['hr.employee'].search([
                            ('name', '=', employee_line),
                        ])
                        employee_job = str(employee_obj.job_id.name)
                for line_dear in doc_ids.dear_select:
                    dear_name.append(str(line_dear))
                    receipt_group = []
                for receipt in doc_ids.reference_receive_document_related:
                    receipt_group.append(receipt)
                attached_file_in = []
                for attached in doc_ids.attached_file_in:
                    attached_file_in.append(attached)

                user_create = self.to_thai_number(str(doc_ids.create_uid.name))
                user_create_obj = self.env['hr.employee'].search([
                    ('user_id', '=', doc_ids.create_uid.id),
                ])

                department_create = self.to_thai_number(str(user_create_obj.department_id.name))
                department_email = self.to_thai_number(str(user_create_obj.work_email))

                department_obj = self.env['hr.department'].search([
                    ('name', '=', doc_ids.department_name),
                ])
                for department_line in department_obj:
                    if department_line.depa_phone:
                        department_phone = self.to_thai_number(str(department_line.depa_phone))
                    if department_line.depa_fax:
                        department_fax = self.to_thai_number(str(department_line.depa_fax))
                    # if department_line.depa_email:
                    #     department_email = self.to_thai_number(str(department_line.depa_email))

                # document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                # last_paragraph = document.paragraphs[-1]
                # last_paragraph.alignment = WD_TABLE_ALIGNMENT.RIGHT

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                # edit_record3 = document.add_paragraph('')
                # edit_record3.add_run('')
                # edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # edit_record3.paragraph_format.space_after = Pt(0)
                # edit_record4 = document.add_paragraph('')
                # edit_record4.add_run('')
                # edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # edit_record4.paragraph_format.space_after = Pt(0)

                zero_record = document.add_paragraph('')
                font = zero_record.add_run(speed).font
                font.size = Pt(24)
                font.color.rgb = RGBColor(255, 0, 0)
                font.bold = True
                # zero_record.style = document.styles['depa_speed']
                # font51 = zero_record.add_run()
                # font51.add_text(speed)
                zero_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                zero_record.paragraph_format.space_after = Pt(0)

                one_record = document.add_paragraph('ที่   ')
                one_record.add_run(real_name)
                one_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                one_record.paragraph_format.space_after = Pt(0)

                one_records = document.add_paragraph('')
                one_records.add_run('')
                one_records.add_run(day)
                one_records.add_run(' ')
                one_records.add_run(month_th)
                one_records.add_run(' ')
                one_records.add_run(year)
                one_records.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                one_records.paragraph_format.space_after = Pt(0)
                one_records.paragraph_format.space_before = Pt(6)
                one_records.paragraph_format.left_indent = Inches(3.149606)

                two_record = document.add_paragraph()
                # two_record.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
                two_record.add_run('เรื่อง').bold = True
                two_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                two_record.paragraph_format.space_after = Pt(0)
                two_record.paragraph_format.space_before = Pt(6)
                two_record.add_run("\t")
                two_record.add_run(subject)
                two_record.paragraph_format.left_indent = Inches(0.4)
                two_record.paragraph_format.first_line_indent = Inches(-0.4)


                three_record = document.add_paragraph()
                # three_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                three_record.add_run('เรียน').bold = True
                three_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                three_record.paragraph_format.space_after = Pt(0)
                three_record.paragraph_format.space_before = Pt(6)
                three_record.add_run("\t")
                three_record.add_run(dear_name)
                three_record.paragraph_format.left_indent = Inches(0.4)
                three_record.paragraph_format.first_line_indent = Inches(-0.4)

                if len(receipt_group) > 0:
                    four_record = document.add_paragraph('อ้างถึง')
                    # four_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    four_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    four_record.paragraph_format.space_after = Pt(0)
                    four_record.paragraph_format.space_before = Pt(0)
                    four_record.add_run("\t")
                    four_record.add_run('หนังสือ')
                    # if receipt_group[0].from_document == False:
                    #     four_record.add_run(' ')
                    if receipt_group[0].from_document != False:
                        four_record.add_run(' '+receipt_group[0].from_document)
                    if receipt_group[0].refer != False:
                        four_record.add_run(' ที่ ')
                        four_record.add_run(self.to_thai_number(str(receipt_group[0].refer)))
                    four_record.add_run(' ลงวันที่ ')
                    four_record.add_run(
                        self.to_thai_number(str(self.to_full_thai_date(receipt_group[0].date_receive))))
                    four_record.paragraph_format.left_indent = Inches(0.5)
                    four_record.paragraph_format.first_line_indent = Inches(-0.5)
                    if len(receipt_group) > 1:
                        four_record1 = document.add_paragraph("\t")
                        four_record1.add_run('หนังสือ')
                        # if receipt_group[1].from_document == False:
                        #     four_record1.add_run(' ')
                        if receipt_group[1].from_document != False:
                            four_record1.add_run(' '+receipt_group[1].from_document)
                        if receipt_group[1].refer != False:
                            four_record1.add_run(' ที่ ')
                            four_record1.add_run(self.to_thai_number(str(receipt_group[1].refer)))
                        four_record1.add_run(' ลงวันที่ ')
                        four_record1.add_run(
                            self.to_thai_number(str(self.to_full_thai_date(receipt_group[1].date_receive))))
                        four_record1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        four_record1.paragraph_format.space_after = Pt(0)
                        four_record1.paragraph_format.space_before = Pt(0)
                        four_record1.paragraph_format.left_indent = Inches(0.5)
                        four_record1.paragraph_format.first_line_indent = Inches(-0.5)
                        if len(receipt_group) > 2:
                            four_record2 = document.add_paragraph("\t")
                            four_record2.add_run('หนังสือ')
                            # if receipt_group[2].from_document == False:
                            #     four_record2.add_run(' ')
                            if receipt_group[2].from_document != False:
                                four_record2.add_run(' '+receipt_group[2].from_document)
                            if receipt_group[2].refer != False:
                                four_record2.add_run(' ที่ ')
                                four_record2.add_run(self.to_thai_number(str(receipt_group[2].refer)))
                            four_record2.add_run(' ลงวันที่ ')
                            four_record2.add_run(
                                self.to_thai_number(str(self.to_full_thai_date(receipt_group[2].date_receive))))
                            four_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            four_record2.paragraph_format.space_after = Pt(0)
                            four_record2.paragraph_format.space_before = Pt(0)
                            four_record2.paragraph_format.left_indent = Inches(0.5)
                            four_record2.paragraph_format.first_line_indent = Inches(-0.5)
                            if len(receipt_group) > 3:
                                four_record3 = document.add_paragraph("\t")
                                four_record3.add_run('หนังสือ')
                                # if receipt_group[3].from_document == False:
                                #     four_record3.add_run(' ')
                                if receipt_group[3].from_document != False:
                                    four_record3.add_run(' '+receipt_group[3].from_document)
                                if receipt_group[3].refer != False:
                                    four_record3.add_run(' ที่ ')
                                    four_record3.add_run(self.to_thai_number(str(receipt_group[3].refer)))
                                four_record3.add_run(' ลงวันที่ ')
                                four_record3.add_run(
                                    self.to_thai_number(str(self.to_full_thai_date(receipt_group[3].date_receive))))
                                four_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                                four_record3.paragraph_format.space_after = Pt(0)
                                four_record3.paragraph_format.space_before = Pt(0)
                                four_record3.paragraph_format.left_indent = Inches(0.5)
                                four_record3.paragraph_format.first_line_indent = Inches(-0.5)

                if len(attached_file_in) > 0:
                    seq = ''
                    if len(attached_file_in) > 1:
                        seq = self.to_single_thai_number(1) + '. '
                    five_record = document.add_paragraph(f'สิ่งที่ส่งมาด้วย')
                    # five_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    five_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    five_record.paragraph_format.space_after = Pt(0)
                    five_record.paragraph_format.space_before = Pt(6)
                    five_record.add_run("\t")
                    five_record.add_run(f'{str(seq)}{str(attached_file_in[0].name)}')
                    five_record.paragraph_format.left_indent = Inches(0.9)
                    five_record.paragraph_format.first_line_indent = Inches(-0.9)
                    if len(attached_file_in) > 1:
                        for i in range(1, len(attached_file_in)):
                            five_record1 = document.add_paragraph("\t")
                            five_record1.add_run(f'{self.to_single_thai_number(i+1)}. {str(attached_file_in[i].name)}')
                            five_record1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            five_record1.paragraph_format.space_after = Pt(0)
                            five_record1.paragraph_format.space_before = Pt(0)
                            five_record1.paragraph_format.left_indent = Inches(0.9)
                            five_record1.paragraph_format.first_line_indent = Inches(-0.9)
                        #     five_record1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        #     five_record1.paragraph_format.space_after = Pt(0)
                        #     five_record1.paragraph_format.space_before = Pt(6)
                        # if len(attached_file_in) > 2:
                        #     five_record2 = document.add_paragraph('                   ')
                        #     five_record2.add_run(self.to_thai_number(str(attached_file_in[2].name)))
                        #     five_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        #     five_record2.paragraph_format.space_after = Pt(0)
                        #     five_record2.paragraph_format.space_before = Pt(6)
                        #     if len(attached_file_in) > 3:
                        #         five_record3 = document.add_paragraph('                  ')
                        #         five_record3.add_run(self.to_thai_number(str(attached_file_in[3].name)))
                        #         five_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        #         five_record3.paragraph_format.space_after = Pt(0)
                        #         five_record3.paragraph_format.space_before = Pt(6)
                    # five_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    # five_record.paragraph_format.space_after = Pt(0)
                    # five_record.paragraph_format.space_before = Pt(0)


                if "\n" in material.strip():
                    for p in material.split("\n"):
                        if p:
                            six_record = document.add_paragraph('')
                            six_record.add_run(p)
                            # six_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            six_record.paragraph_format.space_after = Pt(0)
                            six_record.paragraph_format.space_before = Pt(6)
                            six_record.paragraph_format.first_line_indent = Inches(0.9842520)
                else:
                    six_record = document.add_paragraph('')
                    six_record.add_run(material)
                    # six_record.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    six_record.paragraph_format.space_after = Pt(0)
                    six_record.paragraph_format.space_before = Pt(6)
                    six_record.paragraph_format.first_line_indent = Inches(0.9842520)

                if for_document != 'other':
                    seven_record = document.add_paragraph('จึงเรียนมา')
                    seven_record.add_run(for_document)
                else:
                    seven_record = document.add_paragraph('')
                    seven_record.add_run(for_document_other)
                seven_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                seven_record.paragraph_format.space_after = Pt(0)
                seven_record.paragraph_format.space_before = Pt(6)
                seven_record.paragraph_format.left_indent = Inches(0.9842520)

                seven_record2 = document.add_paragraph('ขอแสดงความนับถือ')
                seven_record2.add_run('')
                seven_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                seven_record2.paragraph_format.space_after = Pt(0)
                seven_record2.paragraph_format.space_before = Pt(6)
                seven_record2.paragraph_format.left_indent = Inches(3.149606)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record3 = document.add_paragraph('')
                edit_record3.add_run('')
                edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record3.paragraph_format.space_after = Pt(0)
                edit_record4 = document.add_paragraph('')
                edit_record4.add_run('')
                edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record4.paragraph_format.space_after = Pt(0)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(employee_line)
                sevens_record.add_run(' )')
                sevens_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                sevens_record.paragraph_format.space_after = Pt(0)
                sevens_record.paragraph_format.space_before = Pt(6)
                sevens_record.paragraph_format.left_indent = Inches(3)
                # sevens_record.paragraph_format.left_indent = Inches(2.952756)

                eight_record = document.add_paragraph('')
                eight_record.add_run(employee_job)
                eight_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                eight_record.paragraph_format.space_after = Pt(0)
                eight_record.paragraph_format.space_before = Pt(6)
                eight_record.paragraph_format.left_indent = Inches(2.5)
                # eight_record.paragraph_format.left_indent = Inches(2.460630)

                edit_record5 = document.add_paragraph('')
                edit_record5.add_run('')
                edit_record5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record5.paragraph_format.space_after = Pt(0)
                # edit_record6 = document.add_paragraph('')
                # edit_record6.add_run('')
                # edit_record6.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # edit_record6.paragraph_format.space_after = Pt(0)
                # edit_record7 = document.add_paragraph('')
                # edit_record7.add_run('')
                # edit_record7.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # edit_record7.paragraph_format.space_after = Pt(0)

                nine_record = document.add_paragraph('')
                nine_record.add_run(department_create)
                nine_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                nine_record.paragraph_format.space_after = Pt(0)
                nine_record.paragraph_format.space_before = Pt(0)

                ten_record = document.add_paragraph('โทรศัพท์  ')
                # ten_record.add_run(department_phone)
                ten_record.add_run(self.to_thai_number('02-026-2333'))
                ten_record.add_run('   (')
                ten_record.add_run(user_create)
                ten_record.add_run(')')
                ten_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                ten_record.paragraph_format.space_after = Pt(0)
                ten_record.paragraph_format.space_before = Pt(0)

                twelve_record = document.add_paragraph('ไปรษณีย์อิเล็กทรอนิกส์ : ')
                twelve_record.add_run(department_email)
                twelve_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                twelve_record.paragraph_format.space_after = Pt(0)
                twelve_record.paragraph_format.space_before = Pt(0)

                # font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run(name).font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run('  ').font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run(user_name).font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run('  ').font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run(user_job_title).font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run('  ').font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                # font.size = Pt(10)
                # font = footer.paragraphs[0].add_run(current_date).font
                # font.size = Pt(10)
            if doc_ids.document_type == 'ระเบียบ':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.subject:
                    subject = self.to_thai_number(str(doc_ids.subject))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                if doc_ids.material:
                    material = str(doc_ids.material)
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                    day = self.to_thai_number(str(int((doc_ids.date_document).strftime('%d'))))
                    month = (int((doc_ids.date_document).strftime('%m')))
                    if month == 1:
                        month_th = 'มกราคม'
                    if month == 2:
                        month_th = 'กุมภาพันธ์'
                    if month == 3:
                        month_th = 'มีนาคม'
                    if month == 4:
                        month_th = 'เมษายน'
                    if month == 5:
                        month_th = 'พฤษภาคม'
                    if month == 6:
                        month_th = 'มิถุนายน'
                    if month == 7:
                        month_th = 'กรกฎาคม'
                    if month == 8:
                        month_th = 'สิงหาคม'
                    if month == 9:
                        month_th = 'กันยายน'
                    if month == 10:
                        month_th = 'ตุลาคม'
                    if month == 11:
                        month_th = 'พฤศจิกายน'
                    if month == 12:
                        month_th = 'ธันวาคม'
                    year = self.to_thai_number(str(int((doc_ids.date_document).strftime('%Y')) + 543))
                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                    len_real_name = str(len(doc_ids.name_real))
                    print(len_real_name)
                for employee_name in doc_ids.empolyee_name:
                    employee.append(str(employee_name.name))
                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)
                        employee_line = str(doc_line.employee_id.name)
                        employee_job = str(doc_line.job_id_name.name)
                department_obj = self.env['hr.department'].search([
                    ('name', '=', doc_ids.department_name),
                ])
                for department_line in department_obj:
                    if department_line.depa_phone:
                        department_phone = department_line.depa_phone
                    if department_line.depa_fax:
                        department_fax = department_line.depa_fax
                    if department_line.depa_email:
                        department_email = department_line.depa_email
                sing_department = self.to_thai_number(str(doc_ids.sign_department.name))
                sing_employee = self.to_thai_number(str(doc_ids.sign_employee.name))

                document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record = document.add_paragraph('')
                font = one_record.add_run('ระเบียบสำนักงานส่งเสริมเศรษฐกิจดิจิทัล').font
                font.size = Pt(18)
                font.bold = True
                one_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record.paragraph_format.space_after = Pt(0)
                one_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                two_record = document.add_paragraph('')
                font1 = two_record.add_run('ว่าด้วย ').font
                font2 = two_record.add_run(subject).font
                font3 = two_record.add_run(' ').font
                font4 = two_record.add_run(date).font
                font1.size = Pt(18)
                font1.bold = True
                font2.size = Pt(18)
                font2.bold = True
                font3.size = Pt(18)
                font3.bold = True
                font4.size = Pt(18)
                font4.bold = True
                two_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                two_record.paragraph_format.space_after = Pt(0)
                two_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                three_record = document.add_paragraph('')
                three_record.add_run('ฉบับที่ ').bold = True
                if len_real_name == '14':
                    font1 = three_record.add_run(real_name[8]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '15':
                    font1 = three_record.add_run(real_name[8:10]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '16':
                    font1 = three_record.add_run(real_name[8:11]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '17':
                    font1 = three_record.add_run(real_name[8:12]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '18':
                    font1 = three_record.add_run(real_name[8:13]).font
                    font1.size = Pt(18)
                    font1.bold = True
                three_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                three_record.paragraph_format.space_after = Pt(0)
                three_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                four_record = document.add_paragraph('')
                four_record.add_run('พ.ศ. ').bold = True
                if len_real_name == '14':
                    font1 = four_record.add_run(real_name[10:14]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '15':
                    font1 = four_record.add_run(real_name[11:15]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '16':
                    font1 = four_record.add_run(real_name[12:16]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '17':
                    font1 = four_record.add_run(real_name[13:17]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '18':
                    font1 = four_record.add_run(real_name[14:18]).font
                    font1.size = Pt(18)
                    font1.bold = True

                four_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                four_record.paragraph_format.space_after = Pt(0)
                four_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                five_record = document.add_paragraph('____________________________________  ')
                five_record.add_run('')
                five_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                edit_record = document.add_paragraph('')
                edit_record.add_run('')
                edit_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record.paragraph_format.space_after = Pt(0)
                edit_record.paragraph_format.space_before = Pt(12)

                six_record = document.add_paragraph('')
                six_record.add_run(material)
                six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                six_record.paragraph_format.space_after = Pt(0)
                six_record.paragraph_format.space_before = Pt(6)

                # edit_record2 = document.add_paragraph('')
                # edit_record2.add_run('')
                # edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # edit_record2.paragraph_format.space_after = Pt(0)
                # edit_record2.paragraph_format.space_before = Pt(12)

                seven_record = document.add_paragraph('ประกาศ ณ วันที่ ')
                seven_record.add_run(day)
                seven_record.add_run(' ')
                seven_record.add_run('เดือน ')
                seven_record.add_run(month_th)
                seven_record.add_run(' ')
                seven_record.add_run('พ.ศ. ')
                seven_record.add_run(year)
                seven_record.paragraph_format.left_indent = Inches(0.9842520)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record3 = document.add_paragraph('')
                edit_record3.add_run('')
                edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record3.paragraph_format.space_after = Pt(0)
                edit_record4 = document.add_paragraph('')
                edit_record4.add_run('')
                edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record4.paragraph_format.space_after = Pt(0)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(sing_employee)
                sevens_record.add_run(' )')
                sevens_record.paragraph_format.left_indent = Inches(3.149606)

                eight_record = document.add_paragraph('')
                eight_record.add_run(sing_department)
                eight_record.paragraph_format.left_indent = Inches(3.051181)
                edits_record = document.add_paragraph('')
                edits_record.add_run('')

                font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_job_title).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(current_date).font
                font.size = Pt(10)
            if doc_ids.document_type == 'ประกาศ':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                sing_department = self.to_thai_number(str(doc_ids.sign_department.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                speed = self.text_speed(str(doc_ids.speed))
                secret = self.text_secret(str(doc_ids.secret))
                print(secret, speed)
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.subject:
                    subject = self.to_thai_number(str(doc_ids.subject))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                if doc_ids.material:
                    material = str(doc_ids.material)
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                    day = self.to_thai_number(str(int((doc_ids.date_document).strftime('%d'))))
                    month = (int((doc_ids.date_document).strftime('%m')))
                    if month == 1:
                        month_th = 'มกราคม'
                    if month == 2:
                        month_th = 'กุมภาพันธ์'
                    if month == 3:
                        month_th = 'มีนาคม'
                    if month == 4:
                        month_th = 'เมษายน'
                    if month == 5:
                        month_th = 'พฤษภาคม'
                    if month == 6:
                        month_th = 'มิถุนายน'
                    if month == 7:
                        month_th = 'กรกฎาคม'
                    if month == 8:
                        month_th = 'สิงหาคม'
                    if month == 9:
                        month_th = 'กันยายน'
                    if month == 10:
                        month_th = 'ตุลาคม'
                    if month == 11:
                        month_th = 'พฤศจิกายน'
                    if month == 12:
                        month_th = 'ธันวาคม'
                    year = self.to_thai_number(str(int((doc_ids.date_document).strftime('%Y')) + 543))
                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                    len_real_name = str(len(doc_ids.name_real))
                for employee_name in doc_ids.empolyee_name:
                    employee.append(str(employee_name.name))
                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)
                        employee_line = str(doc_line.employee_id.name)
                        employee_job = str(doc_line.job_id_name.name)
                department_obj = self.env['hr.department'].search([
                    ('name', '=', doc_ids.department_name),
                ])
                for department_line in department_obj:
                    if department_line.depa_phone:
                        department_phone = department_line.depa_phone
                    if department_line.depa_fax:
                        department_fax = department_line.depa_fax
                    if department_line.depa_email:
                        department_email = department_line.depa_email

                document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                zero_record = document.add_paragraph('')
                zero_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                zero_record.paragraph_format.space_after = Pt(0)
                font1 = zero_record.add_run('ประกาศ').font
                font1.size = Pt(18)
                font1.bold = True
                zero_record.add_run('')
                zero_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                one_record = document.add_paragraph('')
                font1 = one_record.add_run('สำนักงานส่งเสริมเศรษฐกิจดิจิทัล').font
                font1.size = Pt(18)
                font1.bold = True
                one_record.add_run('')
                one_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                one_record.paragraph_format.space_after = Pt(0)
                one_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                two_record = document.add_paragraph('')
                font = two_record.add_run('เรื่อง ').font
                font2 = two_record.add_run(subject).font
                font3 = two_record.add_run(' ').font
                font4 = two_record.add_run(date).font
                font.size = Pt(18)
                font.bold = True
                font2.size = Pt(18)
                font2.bold = True
                font3.size = Pt(18)
                font3.bold = True
                font4.size = Pt(18)
                font4.bold = True
                two_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                two_record.paragraph_format.space_after = Pt(0)
                two_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                five_record = document.add_paragraph('____________________________________  ')
                five_record.add_run('')
                five_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                edit_record = document.add_paragraph('')
                edit_record.add_run('')
                edit_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record.paragraph_format.space_after = Pt(0)
                edit_record.paragraph_format.space_before = Pt(12)

                six_record = document.add_paragraph('')
                six_record.add_run(material)
                six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                six_record.paragraph_format.space_after = Pt(6)
                six_record.paragraph_format.space_before = Pt(0)

                seven_record = document.add_paragraph('ประกาศ ณ วันที่ ')
                seven_record.add_run(day)
                seven_record.add_run(' ')
                seven_record.add_run('เดือน ')
                seven_record.add_run(month_th)
                seven_record.add_run(' ')
                seven_record.add_run('พ.ศ. ')
                seven_record.add_run(year)
                seven_record.paragraph_format.left_indent = Inches(0.9842520)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record3 = document.add_paragraph('')
                edit_record3.add_run('')
                edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record3.paragraph_format.space_after = Pt(0)
                edit_record4 = document.add_paragraph('')
                edit_record4.add_run('')
                edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record4.paragraph_format.space_after = Pt(0)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(employee_line)
                sevens_record.add_run(' )')
                sevens_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                eight_record = document.add_paragraph('')
                eight_record.add_run(sing_department)
                eight_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                edits_record = document.add_paragraph('')
                edits_record.add_run('')
                nine_record = document.add_paragraph('')
                nine_record.add_run(note)

                font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_job_title).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(current_date).font
                font.size = Pt(10)
            if doc_ids.document_type == 'ข้อบังคับ':
                user = self.env['res.users'].search([
                    ('id', '=', self._uid)
                ])
                user_obj = self.env['hr.employee'].search([
                    ('name', '=', user.name),
                ])
                user_name = self.to_thai_number(str(user_obj.name))
                user_job_title = user_obj.job_title
                name = self.to_thai_number(str(doc_ids.name))
                speed = self.text_speed(str(doc_ids.speed))
                secret = self.text_secret(str(doc_ids.secret))
                current_date = self.to_thai_number(str(self.year_convert(datetime.now() + timedelta(hours=7))))
                if doc_ids.department_id.depa_department:
                    depa_department = str(doc_ids.department_id.depa_department)
                if doc_ids.department_id.depa_phone:
                    depa_phone = self.to_thai_number(str(doc_ids.department_id.depa_phone))
                if doc_ids.subject:
                    subject = self.to_thai_number(str(doc_ids.subject))
                if doc_ids.for_document:
                    for_document = str(doc_ids.for_document)
                if doc_ids.material:
                    material = str(doc_ids.material)
                if doc_ids.note:
                    note = self.to_thai_number(str(doc_ids.note))
                create_name = str(doc_ids.create_uid.partner_id.name)
                obj = self.env['hr.employee'].search([
                    ('name', '=', create_name),
                ])
                if doc_ids.date_document:
                    date = self.to_thai_number(str(self.year_convert2(doc_ids.date_document)))
                    day = self.to_thai_number(str(int((doc_ids.date_document).strftime('%d'))))
                    month = (int((doc_ids.date_document).strftime('%m')))
                    if month == 1:
                        month_th = 'มกราคม'
                    if month == 2:
                        month_th = 'กุมภาพันธ์'
                    if month == 3:
                        month_th = 'มีนาคม'
                    if month == 4:
                        month_th = 'เมษายน'
                    if month == 5:
                        month_th = 'พฤษภาคม'
                    if month == 6:
                        month_th = 'มิถุนายน'
                    if month == 7:
                        month_th = 'กรกฎาคม'
                    if month == 8:
                        month_th = 'สิงหาคม'
                    if month == 9:
                        month_th = 'กันยายน'
                    if month == 10:
                        month_th = 'ตุลาคม'
                    if month == 11:
                        month_th = 'พฤศจิกายน'
                    if month == 12:
                        month_th = 'ธันวาคม'
                    year = self.to_thai_number(str(int((doc_ids.date_document).strftime('%Y')) + 543))
                if doc_ids.name_real:
                    real_name = self.to_thai_number(str(doc_ids.name_real))
                    len_real_name = str(len(doc_ids.name_real))
                for employee_name in doc_ids.empolyee_name:
                    employee.append(str(employee_name.name))
                for doc_line in doc_ids.setting_line_ids:
                    if doc_line.is_active == True:
                        if doc_line.job_id_name.name not in state_line2:
                            state_line2.append(doc_line.job_id_name.name)
                        employee_line = str(doc_line.employee_id.name)
                        employee_job = str(doc_line.job_id_name.name)
                department_obj = self.env['hr.department'].search([
                    ('name', '=', doc_ids.department_name),
                ])
                for department_line in department_obj:
                    if department_line.depa_phone:
                        department_phone = department_line.depa_phone
                    if department_line.depa_fax:
                        department_fax = department_line.depa_fax
                    if department_line.depa_email:
                        department_email = department_line.depa_email

                sing_department = self.to_thai_number(str(doc_ids.sign_department.name))
                sing_employee = self.to_thai_number(str(doc_ids.sign_employee.name))

                document.add_picture(os.path.join(os.path.dirname(__file__), 'logo.png'))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record = document.add_paragraph('')
                font = one_record.add_run('ข้อบังคับสำนักงานส่งเสริมเศรษฐกิจดิจิทัล').font
                font.size = Pt(18)
                font.bold = True
                one_record.add_run('')
                one_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                one_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                one_record.paragraph_format.space_after = Pt(0)

                two_record = document.add_paragraph('')
                font = two_record.add_run('ว่าด้วย ').font
                font2 = two_record.add_run(subject).font
                font3 = two_record.add_run(' ').font
                font4 = two_record.add_run(date).font
                font.size = Pt(18)
                font.bold = True
                font2.size = Pt(18)
                font2.bold = True
                font3.size = Pt(18)
                font3.bold = True
                font4.size = Pt(18)
                font4.bold = True
                two_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                two_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                two_record.paragraph_format.space_after = Pt(0)

                three_record = document.add_paragraph('')
                three_record.add_run('ฉบับที่ ').bold = True
                if len_real_name == '16':
                    font1 = three_record.add_run(real_name[10]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '17':
                    font1 = three_record.add_run(real_name[10:12]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '18':
                    font1 = three_record.add_run(real_name[10:13]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '19':
                    font1 = three_record.add_run(real_name[10:14]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '20':
                    font1 = three_record.add_run(real_name[10:15]).font
                    font1.size = Pt(18)
                    font1.bold = True
                three_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                three_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                three_record.paragraph_format.space_after = Pt(0)

                four_record = document.add_paragraph('')
                four_record.add_run('พ.ศ. ').bold = True
                if len_real_name == '14':
                    font1 = four_record.add_run(real_name[10:14]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '15':
                    font1 = four_record.add_run(real_name[11:15]).font
                    font1.size = Pt(18)
                    font1.bold = True
                if len_real_name == '16':
                    font1 = four_record.add_run(real_name[12:16]).font
                    font1.size = Pt(18)
                    font1.bold = True

                if len_real_name == '17':
                    font1 = four_record.add_run(real_name[13:17]).font
                    font1.size = Pt(18)
                    font1.bold = True

                if len_real_name == '18':
                    font1 = four_record.add_run(real_name[14:18]).font
                    font1.size = Pt(18)
                    font1.bold = True

                four_record.alignment = WD_TABLE_ALIGNMENT.CENTER
                four_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                four_record.paragraph_format.space_after = Pt(0)

                five_record = document.add_paragraph('____________________________________  ')
                five_record.add_run('')
                five_record.alignment = WD_TABLE_ALIGNMENT.CENTER

                edit_record = document.add_paragraph('')
                edit_record.add_run('')
                edit_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record.paragraph_format.space_after = Pt(0)
                edit_record.paragraph_format.space_before = Pt(12)

                six_record = document.add_paragraph('')
                six_record.add_run(material)
                six_record.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                six_record.paragraph_format.space_after = Pt(6)
                six_record.paragraph_format.space_before = Pt(0)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record2.paragraph_format.space_before = Pt(12)

                seven_record = document.add_paragraph('ประกาศ ณ วันที่ ')
                seven_record.add_run(day)
                seven_record.add_run(' ')
                seven_record.add_run('เดือน ')
                seven_record.add_run(month_th)
                seven_record.add_run(' ')
                seven_record.add_run('พ.ศ. ')
                seven_record.add_run(year)
                seven_record.paragraph_format.left_indent = Inches(0.9842520)

                edit_record2 = document.add_paragraph('')
                edit_record2.add_run('')
                edit_record2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record2.paragraph_format.space_after = Pt(0)
                edit_record3 = document.add_paragraph('')
                edit_record3.add_run('')
                edit_record3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record3.paragraph_format.space_after = Pt(0)
                edit_record4 = document.add_paragraph('')
                edit_record4.add_run('')
                edit_record4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                edit_record4.paragraph_format.space_after = Pt(0)

                sevens_record = document.add_paragraph('( ')
                sevens_record.add_run(sing_employee)
                sevens_record.add_run(' )')
                sevens_record.paragraph_format.left_indent = Inches(3.149606)

                eight_record = document.add_paragraph('')
                eight_record.add_run(sing_department)
                eight_record.paragraph_format.left_indent = Inches(3.051181)
                edits_record = document.add_paragraph('')
                edits_record.add_run('')

                font = footer.paragraphs[0].add_run('เลขที่อ้างอิง ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_name).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(user_job_title).font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('  ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run('วันที่พิมพ์ ').font
                font.size = Pt(10)
                font = footer.paragraphs[0].add_run(current_date).font
                font.size = Pt(10)

        fp = BytesIO()
        document.save(fp)
        fp.seek(0)
        data = fp.read()
        fp.close()
        return data