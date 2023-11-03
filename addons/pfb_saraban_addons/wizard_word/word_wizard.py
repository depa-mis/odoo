# -*- encoding: utf-8 -*-
import time
from datetime import datetime
from dateutil import relativedelta
from odoo import models, fields, api, _
import base64
from io import BytesIO
from odoo.exceptions import UserError, ValidationError
from odoo.tools import DEFAULT_SERVER_DATE_FORMAT
from docx import Document
from docx.shared import Inches

class WordWizardReport(models.TransientModel):
    _name = 'word.wizard.report'
    _description = 'Export to word'

    attachment_id = fields.Many2one('ir.attachment', 'Attachment')
    
    
    def from_data(self,current_id):
        rec_ids = self.env['document.internal.main'].search([('id','in',current_id)])
        document = Document()
        document.add_heading('หนังสือภายใน',9)
        document.add_paragraph('')
        style = document.styles['Normal']
        font = style.font
        font.name = 'Sarabun'
        for doc_ids in rec_ids:
            circular_letter = str(doc_ids.circular_letter)
            name = str(doc_ids.name)
            date = str(doc_ids.date_document)
            real_name = str(doc_ids.name_real)
            real_date = str(doc_ids.date_document_real)
            depart = str(doc_ids.department_name)
            speed = str(doc_ids.speed)
            secret = str(doc_ids.secret)
            subject = str(doc_ids.subject)
            dear = str(doc_ids.dear)
            employee = str(doc_ids.empolyee_name.name)
            sign = str(doc_ids.sign.name)
            for_document = str(doc_ids.for_document)
            other = str(doc_ids.for_document_other)
            note = str(doc_ids.note)
            material = str(doc_ids.material)
            
            fist_para = document.add_paragraph('หนังสือเวียน ')
            fist_para.add_run(circular_letter)
            sec_para = document.add_paragraph('Name ')
            sec_para.add_run(name)
            thir_para = document.add_paragraph('Date ')
            thir_para.add_run(date)
            four_para = document.add_paragraph('Real Name ')
            four_para.add_run(real_name)
            fift_para = document.add_paragraph('Real Date ')
            fift_para.add_run(real_date)
            six_para = document.add_paragraph('Department ')
            six_para.add_run(depart)
            sev_para = document.add_paragraph('Speed ')
            sev_para.add_run(speed)
            eigh_para = document.add_paragraph('Secret ')
            eigh_para.add_run(secret)
            nine_para = document.add_paragraph('Subject ')
            nine_para.add_run(subject )
            ten_para = document.add_paragraph('Dear ')
            ten_para.add_run(dear)
            sixteen_para = document.add_paragraph('Employee ')
            sixteen_para.add_run(employee)
            ele_para = document.add_paragraph('Sign ')
            ele_para.add_run(sign)
            twel_para = document.add_paragraph('For Document ')
            twel_para.add_run(for_document)
            thirteen_para = document.add_paragraph('Material ')
            thirteen_para.add_run(material)
            fourteen_para = document.add_paragraph('Other ')
            fourteen_para.add_run(other)
            fifteen_para = document.add_paragraph('Note ')
            fifteen_para.add_run(note)
            document.add_paragraph('')

#         document.save('test.docx')
        fp = BytesIO()
        document.save(fp)
        fp.seek(0)
        data = fp.read()
        fp.close()
        return data
    
    @api.multi
    def print_xls_report(self):
        current_id = self._context.get('active_ids')
        data = base64.encodestring(self.from_data(current_id))
        attach_vals = {
            'name': '%s.docx' % ('Document Word Report'),
            'datas': data,
            'datas_fname': '%s.docx' % ('Document Word Report'),
        }
        
        doc_id = self.env['ir.attachment'].create(attach_vals)


        if self.attachment_id:
            try:
                self.attachment_id.unlink()
            except:
                pass
        self.write({'attachment_id': doc_id.id})
        return {
            'type': 'ir.actions.act_url',
            'url': 'web/content/%s?download=true' % (doc_id.id),
            'target': 'current',
        }
        

